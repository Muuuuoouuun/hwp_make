from __future__ import annotations

import re
from pathlib import Path
from copy import deepcopy
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from . import storage
from .exam_templates import (
    ANSWER_SHEET_TITLE,
    ExamTemplate,
    answer_blank_text,
    explanation_entries,
    format_answer,
    get_template,
    needs_answer_blank,
    quick_answer_lines,
    resolve_export_title,
)
from .math_text import normalize_math_token, split_math_text, strip_math_delimiters

CIRCLED_NUMBERS = ("①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨")
QUESTION_PREFIX_RE = re.compile(r"^\s*(?:문제\s*)?(\d{1,3})\s*[\.\)]\s*")
CHOICE_PREFIX_RE = re.compile(r"^\s*(?:[①②③④⑤⑥⑦⑧⑨]|\d+\s*[\.\)]|[1-9](?=\s))\s*")

# XML에서 허용되지 않는 제어문자. python-docx는 이런 문자가 있으면 저장 시 예외를 던진다.
_XML_ILLEGAL_RE = re.compile("[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff￾￿]")


def _clean(value: Any) -> str:
    return _XML_ILLEGAL_RE.sub("", str(value or ""))


def _sanitize_problem(problem: dict[str, Any]) -> dict[str, Any]:
    clean = dict(problem)
    for key in ("number", "title", "stem", "answer", "explanation", "subject", "unit"):
        if key in clean:
            clean[key] = _clean(clean[key])
    clean["choices"] = [_clean(choice) for choice in (problem.get("choices") or [])]
    clean["tables"] = [
        [[_clean(cell) for cell in row] for row in table]
        for table in (problem.get("tables") or [])
    ]
    return clean


def _style_run(run, size: int = 10, bold: bool = False, math: bool = False) -> None:
    font_name = "Cambria Math" if math else "Malgun Gothic"
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name if math else "Malgun Gothic")


_COMMAND_TEXT = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "theta": "θ",
    "lambda": "λ",
    "mu": "μ",
    "pi": "π",
    "sigma": "σ",
    "omega": "ω",
    "Delta": "Δ",
    "nabla": "∇",
    "leq": "≤",
    "geq": "≥",
    "neq": "≠",
    "approx": "≈",
    "cdot": "·",
    "times": "×",
    "div": "÷",
    "pm": "±",
    "mp": "∓",
    "infty": "∞",
    "angle": "∠",
    "triangle": "△",
    "parallel": "∥",
    "perp": "⊥",
    "because": "∵",
    "therefore": "∴",
    "sin": "sin",
    "cos": "cos",
    "tan": "tan",
    "sec": "sec",
    "csc": "csc",
    "cot": "cot",
    "log": "log",
    "ln": "ln",
    "lim": "lim",
}
_NARY_SYMBOLS = {"sum": "∑", "prod": "∏", "int": "∫", "iint": "∫∫"}
_ACCENT_CHARS = {
    "overline": "\u0305",
    "bar": "\u0305",
    "vec": "\u20d7",
    "underline": "\u0332",
}


def _m_element(name: str) -> Any:
    return OxmlElement(f"m:{name}")


def _m_run(text: str) -> Any:
    run = _m_element("r")
    text_node = _m_element("t")
    text_node.text = str(text or "")
    run.append(text_node)
    return run


def _append_math_arg(parent: Any, name: str, source: str | None = None, children: list[Any] | None = None) -> None:
    node = _m_element(name)
    payload = children if children is not None else _math_children_from_source(source or "")
    if not payload:
        payload = [_m_run("")]
    for child in payload:
        node.append(child)
    parent.append(node)


def _read_command(source: str, index: int) -> tuple[str, int] | None:
    if index >= len(source) or source[index] != "\\":
        return None
    cursor = index + 1
    while cursor < len(source) and source[cursor].isalpha():
        cursor += 1
    if cursor == index + 1:
        return source[index + 1 : index + 2], min(len(source), index + 2)
    return source[index + 1 : cursor], cursor


def _read_braced(source: str, index: int) -> tuple[str, int] | None:
    if index >= len(source) or source[index] != "{":
        return None
    depth = 0
    start = index + 1
    cursor = index
    while cursor < len(source):
        char = source[cursor]
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return source[start:cursor], cursor + 1
        cursor += 1
    return None


def _read_bracketed(source: str, index: int) -> tuple[str, int] | None:
    if index >= len(source) or source[index] != "[":
        return None
    end = source.find("]", index + 1)
    if end < 0:
        return None
    return source[index + 1 : end], end + 1


def _read_script_arg(source: str, index: int) -> tuple[str, int]:
    braced = _read_braced(source, index)
    if braced is not None:
        return braced
    command = _read_command(source, index)
    if command is not None:
        name, end = command
        return _COMMAND_TEXT.get(name, "\\" + name), end
    return source[index : index + 1], min(len(source), index + 1)


def _read_plain_atom(source: str, index: int) -> tuple[str, int]:
    command = _read_command(source, index)
    if command is not None:
        name, end = command
        return _COMMAND_TEXT.get(name, "\\" + name), end

    cursor = index
    while cursor < len(source):
        char = source[cursor]
        if char in "\\{}_^" or char.isspace():
            break
        cursor += 1
    if cursor == index:
        return source[index : index + 1], index + 1
    return source[index:cursor], cursor


def _skip_spaces(source: str, index: int) -> int:
    cursor = index
    while cursor < len(source) and source[cursor].isspace():
        cursor += 1
    return cursor


def _read_wrapped_source(source: str, index: int) -> tuple[str, int] | None:
    pairs = {"(": ")", "[": "]"}
    opener = source[index : index + 1]
    closer = pairs.get(opener)
    if closer is None:
        return None
    depth = 0
    cursor = index
    while cursor < len(source):
        char = source[cursor]
        if char == opener:
            depth += 1
        elif char == closer:
            depth -= 1
            if depth == 0:
                return source[index : cursor + 1], cursor + 1
        cursor += 1
    return None


def _read_decorated_arg(source: str, index: int) -> tuple[str, int] | None:
    cursor = _skip_spaces(source, index)
    braced = _read_braced(source, cursor)
    if braced is not None:
        return braced
    return _read_atom_source(source, cursor)


def _read_command_source(source: str, index: int) -> tuple[str, int] | None:
    command = _read_command(source, index)
    if command is None:
        return None
    name, next_cursor = command
    if name in {"frac", "dfrac", "tfrac"}:
        numerator = _read_braced(source, next_cursor)
        if numerator is not None:
            denominator = _read_braced(source, numerator[1])
            if denominator is not None:
                return source[index : denominator[1]], denominator[1]
    if name == "sqrt":
        cursor = next_cursor
        bracketed = _read_bracketed(source, cursor)
        if bracketed is not None:
            cursor = bracketed[1]
        base = _read_braced(source, cursor)
        if base is not None:
            return source[index : base[1]], base[1]
    if name in _ACCENT_CHARS:
        base = _read_decorated_arg(source, next_cursor)
        if base is not None:
            return source[index : base[1]], base[1]
    if name in _NARY_SYMBOLS:
        _, _, cursor = _scripts_after(source, next_cursor)
        body = _read_atom_source(source, cursor)
        return (source[index : body[1]], body[1]) if body is not None else (source[index:cursor], cursor)
    _, _, cursor = _scripts_after(source, next_cursor)
    return source[index:cursor], cursor


def _read_atom_source(source: str, index: int) -> tuple[str, int] | None:
    cursor = _skip_spaces(source, index)
    if cursor >= len(source) or source[cursor] in "+-=*/<>^_":
        return None
    braced = _read_braced(source, cursor)
    if braced is not None:
        return braced
    wrapped = _read_wrapped_source(source, cursor)
    if wrapped is not None:
        return wrapped
    command = _read_command_source(source, cursor)
    if command is not None:
        return command
    atom_text, next_cursor = _read_plain_atom(source, cursor)
    if not atom_text:
        return None
    _, _, end = _scripts_after(source, next_cursor)
    return source[cursor:end], end


def _omml_fraction(numerator: str, denominator: str) -> Any:
    node = _m_element("f")
    _append_math_arg(node, "num", numerator)
    _append_math_arg(node, "den", denominator)
    return node


def _omml_radical(base: str, degree: str = "") -> Any:
    node = _m_element("rad")
    if degree:
        _append_math_arg(node, "deg", degree)
    _append_math_arg(node, "e", base)
    return node


def _omml_script(base_children: list[Any], sub: str, sup: str) -> Any:
    if sub and sup:
        node = _m_element("sSubSup")
        _append_math_arg(node, "e", children=[deepcopy(child) for child in base_children])
        _append_math_arg(node, "sub", sub)
        _append_math_arg(node, "sup", sup)
        return node
    if sub:
        node = _m_element("sSub")
        _append_math_arg(node, "e", children=[deepcopy(child) for child in base_children])
        _append_math_arg(node, "sub", sub)
        return node
    node = _m_element("sSup")
    _append_math_arg(node, "e", children=[deepcopy(child) for child in base_children])
    _append_math_arg(node, "sup", sup)
    return node


def _omml_accent(accent: str, base: str) -> Any:
    node = _m_element("acc")
    props = _m_element("accPr")
    chr_node = _m_element("chr")
    chr_node.set(qn("m:val"), accent)
    props.append(chr_node)
    node.append(props)
    _append_math_arg(node, "e", base)
    return node


def _omml_delimiter(begin: str, end: str, body: str | None = None, children: list[Any] | None = None) -> Any:
    node = _m_element("d")
    props = _m_element("dPr")
    beg_node = _m_element("begChr")
    beg_node.set(qn("m:val"), begin or ".")
    end_node = _m_element("endChr")
    end_node.set(qn("m:val"), end or ".")
    props.append(beg_node)
    props.append(end_node)
    node.append(props)
    _append_math_arg(node, "e", body or "", children=children)
    return node


def _omml_matrix(rows: list[list[str]]) -> Any | None:
    clean_rows = [[cell.strip() for cell in row] for row in rows if any(cell.strip() for cell in row)]
    if not clean_rows:
        return None
    node = _m_element("m")
    for row in clean_rows:
        row_node = _m_element("mr")
        for cell in row:
            _append_math_arg(row_node, "e", cell)
        node.append(row_node)
    return node


def _omml_nary(symbol: str, sub: str = "", sup: str = "", body: str = "") -> Any:
    node = _m_element("nary")
    props = _m_element("naryPr")
    chr_node = _m_element("chr")
    chr_node.set(qn("m:val"), symbol)
    props.append(chr_node)
    node.append(props)
    if sub:
        _append_math_arg(node, "sub", sub)
    if sup:
        _append_math_arg(node, "sup", sup)
    _append_math_arg(node, "e", body)
    return node


def _scripts_after(source: str, index: int) -> tuple[str, str, int]:
    sub = ""
    sup = ""
    cursor = index
    while cursor < len(source) and source[cursor] in "_^":
        marker = source[cursor]
        value, cursor = _read_script_arg(source, cursor + 1)
        if marker == "_":
            sub = value
        else:
            sup = value
    return sub, sup, cursor


def _read_latex_delimiter(source: str, index: int) -> tuple[str, int] | None:
    cursor = _skip_spaces(source, index)
    if cursor >= len(source):
        return None
    named = {
        r"\{": "{",
        r"\}": "}",
        r"\langle": "〈",
        r"\rangle": "〉",
        r"\|": "‖",
    }
    for token, value in named.items():
        if source.startswith(token, cursor):
            return value, cursor + len(token)
    if source[cursor] == "\\" and cursor + 1 < len(source):
        return source[cursor + 1], cursor + 2
    return source[cursor], cursor + 1


def _read_left_right(source: str, index: int) -> tuple[str, str, str, int] | None:
    command = _read_command(source, index)
    if command is None or command[0] != "left":
        return None
    begin = _read_latex_delimiter(source, command[1])
    if begin is None:
        return None
    begin_token, body_start = begin
    depth = 1
    cursor = body_start
    while cursor < len(source):
        nested = _read_command(source, cursor)
        if nested is not None and nested[0] == "left":
            depth += 1
            cursor = nested[1]
            continue
        if nested is not None and nested[0] == "right":
            depth -= 1
            if depth == 0:
                end = _read_latex_delimiter(source, nested[1])
                if end is None:
                    return None
                end_token, end_cursor = end
                return begin_token, source[body_start:cursor].strip(), end_token, end_cursor
            cursor = nested[1]
            continue
        cursor += 1
    return None


def _read_begin_environment(source: str, index: int) -> tuple[str, str, int] | None:
    command = _read_command(source, index)
    if command is None or command[0] != "begin":
        return None
    env = _read_braced(source, command[1])
    if env is None:
        return None
    env_name, body_start = env
    end_marker = rf"\end{{{env_name}}}"
    end_index = source.find(end_marker, body_start)
    if end_index < 0:
        return None
    return env_name, source[body_start:end_index], end_index + len(end_marker)


def _latex_matrix_rows(body: str) -> list[list[str]]:
    rows = []
    for row in re.split(r"\\\\", body):
        cells = [cell.strip() for cell in row.split("&")]
        if any(cells):
            rows.append(cells)
    return rows


def _math_children_from_source(source: str) -> list[Any]:
    expr = normalize_math_token(strip_math_delimiters(source)).strip()
    children: list[Any] = []
    cursor = 0
    while cursor < len(expr):
        char = expr[cursor]
        if char.isspace():
            cursor += 1
            continue

        if char == "√":
            base = _read_atom_source(expr, cursor + 1)
            if base is not None:
                body, cursor = base
                children.append(_omml_radical(body))
                continue

        if char in {"∑", "∏", "∫"}:
            sub, sup, cursor = _scripts_after(expr, cursor + 1)
            body = _read_atom_source(expr, cursor)
            body_text = ""
            if body is not None:
                body_text, cursor = body
            children.append(_omml_nary(char, sub, sup, body_text))
            continue

        command = _read_command(expr, cursor)
        if command is not None:
            name, next_cursor = command
            if name == "left":
                group = _read_left_right(expr, cursor)
                if group is not None:
                    begin, body, end, cursor = group
                    children.append(_omml_delimiter(begin, end, body))
                    continue
            if name == "begin":
                environment = _read_begin_environment(expr, cursor)
                if environment is not None:
                    env_name, body, cursor = environment
                    matrix = _omml_matrix(_latex_matrix_rows(body))
                    if matrix is not None:
                        if env_name == "pmatrix":
                            children.append(_omml_delimiter("(", ")", children=[matrix]))
                        elif env_name == "bmatrix":
                            children.append(_omml_delimiter("[", "]", children=[matrix]))
                        elif env_name == "cases":
                            children.append(_omml_delimiter("{", ".", children=[matrix]))
                        else:
                            children.append(matrix)
                        continue
            if name in {"frac", "dfrac", "tfrac"}:
                numerator = _read_braced(expr, next_cursor)
                if numerator is not None:
                    denominator = _read_braced(expr, numerator[1])
                    if denominator is not None:
                        children.append(_omml_fraction(numerator[0], denominator[0]))
                        cursor = denominator[1]
                        continue
            if name == "sqrt":
                degree = ""
                next_after_degree = next_cursor
                bracketed = _read_bracketed(expr, next_cursor)
                if bracketed is not None:
                    degree, next_after_degree = bracketed
                base = _read_braced(expr, next_after_degree)
                if base is not None:
                    children.append(_omml_radical(base[0], degree))
                    cursor = base[1]
                    continue
            if name in _NARY_SYMBOLS:
                sub, sup, cursor = _scripts_after(expr, next_cursor)
                body = _read_atom_source(expr, cursor)
                body_text = ""
                if body is not None:
                    body_text, cursor = body
                children.append(_omml_nary(_NARY_SYMBOLS[name], sub, sup, body_text))
                continue
            if name in _ACCENT_CHARS:
                base = _read_decorated_arg(expr, next_cursor)
                if base is not None:
                    children.append(_omml_accent(_ACCENT_CHARS[name], base[0]))
                    cursor = base[1]
                    continue
            if name in _COMMAND_TEXT:
                atom = [_m_run(_COMMAND_TEXT[name])]
                sub, sup, cursor = _scripts_after(expr, next_cursor)
                children.append(_omml_script(atom, sub, sup) if sub or sup else atom[0])
                continue

        atom_text, next_cursor = _read_plain_atom(expr, cursor)
        if not atom_text:
            cursor = next_cursor
            continue
        atom = [_m_run(atom_text)]
        sub, sup, cursor = _scripts_after(expr, next_cursor)
        children.append(_omml_script(atom, sub, sup) if sub or sup else atom[0])
    return children


def _omml_math(token: str) -> Any | None:
    children = _math_children_from_source(token)
    if not children:
        return None
    node = _m_element("oMath")
    for child in children:
        node.append(child)
    return node


def _add_omml_math(paragraph: Any, token: str) -> bool:
    try:
        node = _omml_math(token)
    except Exception:
        return False
    if node is None:
        return False
    paragraph._p.append(node)
    return True


def _set_font(paragraph, size: int = 10, bold: bool = False) -> None:
    for run in paragraph.runs:
        _style_run(run, size, bold)


def _add_text_runs(paragraph, text: str, size: int = 10, bold: bool = False) -> None:
    parts = split_math_text(text)
    if not parts:
        run = paragraph.add_run("")
        _style_run(run, size, bold)
        return
    for segment, is_math in parts:
        if is_math and _add_omml_math(paragraph, segment):
            continue
        run = paragraph.add_run(segment)
        _style_run(run, size, bold, math=is_math)


def _add_text_paragraph(document: Document, text: str = "", size: int = 10, bold: bool = False):
    paragraph = document.add_paragraph()
    _add_text_runs(paragraph, text, size=size, bold=bold)
    return paragraph


def _choice_label(index: int, template: ExamTemplate) -> str:
    if template.choice_style == "bare_number":
        return str(index)
    if template.circled_choices and index <= len(CIRCLED_NUMBERS):
        return CIRCLED_NUMBERS[index - 1]
    return f"{index})"


def _strip_question_prefix(text: str, label: str) -> str:
    match = QUESTION_PREFIX_RE.match(text)
    if match and match.group(1) == str(label):
        return text[match.end() :].lstrip()
    return text


def _format_choice(index: int, choice: str, template: ExamTemplate) -> str:
    clean = CHOICE_PREFIX_RE.sub("", choice or "").strip()
    return f"{_choice_label(index, template)} {clean}".rstrip()


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows or not any(rows):
        return
    col_cnt = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_cnt)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(col_cnt):
            cell = table.rows[r].cells[c]
            paragraph = cell.paragraphs[0]
            _add_text_runs(paragraph, str(row[c]) if c < len(row) else "")


def _add_masthead(document: Document, title: str, template: ExamTemplate) -> None:
    heading = document.add_heading(template.masthead_title or title, level=1)
    _set_font(heading, 16, True)

    meta = "   ".join(
        part for part in (template.area, template.period, template.variant) if part
    )
    if meta:
        _add_text_paragraph(document, meta, 11, True)
    if template.show_student_fields:
        _add_text_paragraph(
            document,
            "성명 ____________     수험 번호 ____________     " + template.selection,
            9,
            False,
        )
    elif template.selection:
        _add_text_paragraph(document, template.selection, 9, False)
    for direction in template.directions:
        _add_text_paragraph(document, direction, 9, False)


def _add_answer_sheet(document: Document, problems: list[dict[str, Any]], template: ExamTemplate) -> None:
    document.add_page_break()
    heading = document.add_heading(ANSWER_SHEET_TITLE, level=1)
    _set_font(heading, 14, True)

    _add_text_paragraph(document, "빠른 정답", 11, True)
    for line in quick_answer_lines(problems, template):
        _add_text_paragraph(document, line)

    entries = explanation_entries(problems, template)
    if not entries:
        return
    document.add_paragraph("")
    _add_text_paragraph(document, "해설", 11, True)
    for entry_heading, lines in entries:
        _add_text_paragraph(document, entry_heading, 10, True)
        for line in lines:
            _add_text_paragraph(document, line)


def _set_section_columns(document: Document, columns: int) -> None:
    if columns <= 1:
        return
    section = document.sections[0]
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(columns))
    cols.set(qn("w:space"), "567")


def _docx_image_width(template: ExamTemplate) -> Cm:
    return Cm(7.0 if template.columns > 1 else 14.5)


def write_docx(
    path: Path,
    title: str,
    problems: list[dict[str, Any]],
    template_key: str = "basic",
    include_answer_sheet: bool = False,
) -> None:
    template = get_template(template_key)
    title = _clean(resolve_export_title(title, template))
    problems = [_sanitize_problem(problem) for problem in problems]
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"].font.size = Pt(10)
    _set_section_columns(document, max(1, min(template.columns, 2)))

    _add_masthead(document, title, template)

    for index, problem in enumerate(problems, start=1):
        label = problem.get("number") or str(index)
        subject = problem.get("subject") or ""
        unit = problem.get("unit") or ""
        meta = " / ".join(part for part in [subject, unit] if part)

        stem = problem.get("stem") or ""
        stem_lines = stem.splitlines()
        if template.merge_question_number:
            first_line = (
                _strip_question_prefix(stem_lines[0], label)
                if stem_lines
                else problem.get("title") or "문제"
            )
            paragraph = document.add_paragraph()
            _add_text_runs(paragraph, f"{label}. {first_line}", 11, True)
            if meta:
                _add_text_runs(paragraph, f"  [{meta}]", 9, False)
            for line in stem_lines[1:]:
                _add_text_paragraph(document, line)
        else:
            paragraph = document.add_paragraph()
            _add_text_runs(paragraph, f"{label}. {problem.get('title') or '문제'}", 11, True)
            if meta:
                _add_text_runs(paragraph, f"  [{meta}]", 9, False)
            if stem:
                for line in stem_lines:
                    _add_text_paragraph(document, line)

        for table_rows in problem.get("tables") or []:
            _add_table(document, table_rows)

        for image_path in problem.get("image_paths") or []:
            full_path = storage.resolve_data_image_path(image_path)
            if full_path is None:
                continue
            try:
                document.add_picture(str(full_path), width=_docx_image_width(template))
            except Exception:
                _add_text_paragraph(document, f"[이미지 삽입 실패: {Path(str(image_path)).name}]")

        choices = [
            _format_choice(choice_index, choice, template)
            for choice_index, choice in enumerate(problem.get("choices") or [], start=1)
        ]
        if choices and template.inline_short_choices and sum(len(choice) for choice in choices) <= 90:
            _add_text_paragraph(document, "    ".join(choices))
        else:
            for choice in choices:
                _add_text_paragraph(document, choice)
        if needs_answer_blank(problem, template):
            _add_text_paragraph(document, answer_blank_text(template))

        answer = problem.get("answer") or ""
        explanation = problem.get("explanation") or ""
        if template.include_answers and answer:
            _add_text_paragraph(document, f"정답: {format_answer(problem, template)}")
        if template.include_explanations and explanation:
            _add_text_paragraph(document, f"해설: {explanation}")
        if index != len(problems):
            document.add_paragraph("")

    if include_answer_sheet:
        _add_answer_sheet(document, problems, template)

    document.save(path)
