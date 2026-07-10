from __future__ import annotations

import base64
import csv
import io
import re
import shutil
import sqlite3
import struct
import uuid
import zipfile
import zlib
from datetime import datetime
from pathlib import Path
from typing import Any

from lxml import etree
from PIL import Image
from pypdf import PdfReader

from . import storage
from .math_text import normalize_recognized_math_text

try:  # OCR은 선택 사항: pytesseract + tesseract 실행 파일이 있을 때만 사용
    import pytesseract

    pytesseract.get_tesseract_version()
    HAS_OCR = True
except Exception:
    pytesseract = None
    HAS_OCR = False

try:  # rhwp(러스트 HWP 엔진)는 선택 사항: 있으면 HWP 텍스트 추출에 우선 사용
    import rhwp
except Exception:
    rhwp = None


SAFE_NAME_RE = re.compile(r"[^0-9A-Za-z가-힣._ -]+")
QUESTION_START_RE = re.compile(
    r"(?m)(?=^\s*(?:문제\s*)?\d{1,3}\s*[\.\)]|\n\s*(?:문제\s*)?\d{1,3}\s*[\.\)])"
)
CIRCLED_CHOICE_MARKERS = "①②③④⑤⑥⑦⑧⑨"
INLINE_CIRCLED_CHOICE_RE = re.compile(rf"([{CIRCLED_CHOICE_MARKERS}])\s*")
INLINE_NUMERIC_CHOICE_RE = re.compile(r"(?<![\w/])([1-5])(?:[\.\)])?(?:\s+|$)")
CHOICE_LINE_RE = re.compile(
    rf"^\s*(?P<label>[{CIRCLED_CHOICE_MARKERS}]|\(?[1-9]\)|[1-9][\.\)]|[A-Ea-e][\.\)])\s*(?P<body>.+)$"
)
BARE_NUMERIC_CHOICE_LINE_RE = re.compile(r"^\s*(?P<label>[1-5])\s+(?P<body>\S.+)$")
NUMERIC_CHOICE_LABEL_RE = re.compile(r"^\(?([1-9])[\.\)]?$")
PDF_CHOICE_FRACTION_LABELS = "①②③④⑤"
PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS = "\u25a1\u25a2\ue06d"
PDF_CHOICE_FRACTION_RE = re.compile(
    rf"^\s*(?P<label>[{PDF_CHOICE_FRACTION_LABELS}])\s*(?P<sign>-?)"
    rf"\s*[{PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS}](?P<den>[A-Za-z0-9α-ωΑ-Ωπ∞+\-.]+)?\s*$"
)
PDF_CHOICE_EXPONENT_FRACTION_RE = re.compile(
    rf"^\s*(?P<label>[{PDF_CHOICE_FRACTION_LABELS}])\s*"
    rf"(?P<base>[A-Za-z0-9α-ωΑ-Ωπ∞+\-.]+)"
    rf"\s*[{PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS}](?P<den>[A-Za-z0-9α-ωΑ-Ωπ∞+\-.]+)?\s*$"
)
PDF_CHOICE_UNIT_FRACTION_RE = re.compile(
    rf"^\s*(?P<sign>-?)\s*[{PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS}](?P<den>[0-9]{{1,3}})\s*$"
)
PDF_CHOICE_FRACTION_PART_RE = re.compile(
    r"^[A-Za-z0-9α-ωΑ-Ωπ∞√∑∫+\-./*(){}\[\]\\_=^]+$"
)
PDF_CHOICE_GEOMETRY_LABEL_RE = re.compile(
    rf"^\s*(?P<label>[{PDF_CHOICE_FRACTION_LABELS}])\s*(?P<body>.*)$"
)
PDF_STEM_FRACTION_PART_RE = re.compile(
    r"^[A-Za-z0-9α-ωΑ-Ωπ∞√∑∫+\-./*(){}\[\]\\_^]+$"
)
ANSWER_LINE_RE = re.compile(r"^\s*(?:정답|답)(?:\s*[:：]|\s+)(?P<value>.+?)\s*$")
EXPLANATION_LINE_RE = re.compile(r"^\s*(?:해설|풀이)(?:(?:\s*[:：]|\s+)(?P<value>.*))?\s*$")
SCORE_MARKER_RE = re.compile(r"\[\s*\d+\s*점\s*\]")
PDF_FORM_LABEL_RE = re.compile(r"\(?\s*(?:홀수형|짝수형|미적분|확률과\s*통계|기하)\s*\)?")


class _Sink:
    """가져오기 중 문항을 만들며 중복은 건너뛰고 그 수를 센다."""

    __slots__ = ("created", "skipped", "_seen")

    def __init__(self) -> None:
        self.created: list[dict[str, Any]] = []
        self.skipped = 0
        self._seen: set[str] = set()

    def add(self, data: dict[str, Any]) -> dict[str, Any] | None:
        problem = storage.create_problem_unique(data, self._seen)
        if problem is None:
            self.skipped += 1
        else:
            self.created.append(problem)
        return problem


def _dedup_notices(sink: _Sink) -> list[str]:
    return [f"중복 {sink.skipped}개는 이미 있는 문항이라 건너뛰었습니다."] if sink.skipped else []


def safe_filename(filename: str) -> str:
    name = Path(filename or "upload").name.strip() or "upload"
    name = SAFE_NAME_RE.sub("_", name)
    return name[:120]


def decode_base64(data: str) -> bytes:
    if "," in data and data.split(",", 1)[0].startswith("data:"):
        data = data.split(",", 1)[1]
    return base64.b64decode(data)


def save_upload(filename: str, payload: bytes) -> str:
    storage.ensure_dirs()
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = storage.UPLOAD_DIR / f"{stamp}_{uuid.uuid4().hex[:8]}_{safe_filename(filename)}"
    path.write_bytes(payload)
    return path.relative_to(storage.DATA_DIR).as_posix()


def _clean_text(value: str) -> str:
    value = normalize_recognized_math_text(value)
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[ \t]+\n", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _recognized_pdf_loose_duplicate_key(number: str, stem: str) -> str:
    """Fingerprint repeated odd/even-form PDF problems despite choice/image drift."""
    value = normalize_recognized_math_text(str(stem or ""))
    score = SCORE_MARKER_RE.search(value)
    if score:
        value = value[: score.end()]
    else:
        value = "\n".join(value.splitlines()[:6])
    value = re.sub(r"^\s*\d{1,3}\s*[.)]\s*", "", value)
    value = PDF_FORM_LABEL_RE.sub("", value)
    value = re.sub(r"[^0-9A-Za-z가-힣α-ωΑ-Ω]+", "", value)
    if len(value) < 24:
        return ""
    return f"{str(number or '').strip()}:{value[:180]}"


def _looks_like_math_pdf(*values: str) -> bool:
    return any("수학" in str(value or "") or "math" in str(value or "").lower() for value in values)


def _decode_text(payload: bytes) -> str:
    """텍스트/CSV류 입력은 UTF-8을 우선하고, 오래된 윈도우 한글(cp949)도 받는다."""
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


def _split_questions(page_text: str) -> list[str]:
    text = _clean_text(page_text)
    if not text:
        return []
    chunks = [chunk.strip() for chunk in QUESTION_START_RE.split(text) if chunk.strip()]
    if len(chunks) <= 1:
        paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
        if _looks_like_single_passage(text, paragraphs):
            return [text]
        return paragraphs if len(paragraphs) > 1 else [text]
    return chunks


def _looks_like_single_passage(text: str, paragraphs: list[str]) -> bool:
    """번호 없는 국어 지문처럼 긴 글은 빈 줄 기준으로 여러 문항처럼 쪼개지 않는다."""
    if len(paragraphs) <= 1:
        return False
    if re.search(r"\[(?:지문|보기|자료)\]|\<(?:보기|지문)\>", text):
        return True
    long_paragraphs = [para for para in paragraphs if len(para) >= 80]
    return len(text) >= 600 and len(long_paragraphs) >= 2


def _plain_text_chunks(text: str) -> list[dict[str, Any]]:
    """줄글/붙여넣기 텍스트를 문항 chunk로 바꾼다."""
    clean = _clean_text(text)
    if not clean:
        return []
    line_blocks = [(line, [], []) for line in clean.splitlines()]
    chunks = _paragraphs_to_chunks(line_blocks)
    if len(chunks) > 1 or any(chunk.get("number_hint") for chunk in chunks):
        return chunks

    split = [{"text": chunk, "images": [], "tables": []} for chunk in _split_questions(clean)]
    if len(split) > 1 and not QUESTION_LINE_RE.match(split[0]["text"]):
        first_numbered = next(
            (index for index, chunk in enumerate(split[1:], start=1) if QUESTION_LINE_RE.match(chunk["text"])),
            None,
        )
        if first_numbered is not None:
            split[first_numbered]["text"] = f"{split[0]['text']}\n{split[first_numbered]['text']}"
            split.pop(0)
    return split


def _extract_number(text: str, fallback: int) -> str:
    match = re.match(r"\s*(?:문제\s*)?(\d{1,3})\s*[\.\)]", text)
    return match.group(1) if match else str(fallback)


def import_pdf(filename: str, payload: bytes, metadata: dict[str, Any]) -> dict[str, Any]:
    """PDF 가져오기.

    1순위: 인식 파이프라인(app.recognition) — born-digital 시험지를 문항 단위로 분리하고
           텍스트, 수식 PUA, 라인/문자 좌표, 그림 후보를 보존한다. 텍스트 신뢰도가 낮은
           일부 영역만 보존용 이미지 fallback을 붙인다.
    2순위: 스캔/텍스트레이어 없음 등 결정론적 인식이 불가능한 경우 pypdf+OCR 경로.
    """
    recognized = _import_pdf_recognized(filename, payload, metadata)
    if recognized is not None:
        return recognized
    return _legacy_import_pdf(filename, payload, metadata)


def _import_pdf_recognized(
    filename: str, payload: bytes, metadata: dict[str, Any]
) -> dict[str, Any] | None:
    """인식 파이프라인으로 PDF 문항을 가져온다. born-digital 이 아니면 None(→ 레거시)."""
    try:
        from .recognition.pipeline import recognize_pdf
    except Exception:
        return None  # 인식 스택(fitz 등) 불가 → 레거시로
    try:
        result = recognize_pdf(payload, filename=filename)
    except Exception:
        return None
    if not result.found:
        return None  # 스캔/텍스트레이어 없음 → 레거시가 처리

    save_upload(filename, payload)  # 원본 PDF 보관
    stem_name = Path(filename).stem
    sink = _Sink()
    loose_seen: set[str] = set()
    loose_dedup_enabled = _looks_like_math_pdf(filename, getattr(result, "exam_title", ""))

    def layout_payload(prob: Any) -> dict[str, Any]:
        box = getattr(prob, "box", None)
        bbox_px = None
        if box is not None:
            bbox_px = [box.left, box.top, box.width, box.height]
        pdf_lines = list(getattr(prob, "line_geometries", []) or [])
        return {
            "column_count": int(getattr(prob, "column_count", 0) or 0),
            "column_index": int(getattr(prob, "column_index", 0) or 0),
            "page": {
                "number": int(getattr(prob, "page_number", 0) or 0),
                "width_px": int(getattr(prob, "page_width_px", 0) or 0),
                "height_px": int(getattr(prob, "page_height_px", 0) or 0),
            },
            "bbox_px": bbox_px,
            "block_type": "image_fallback" if getattr(prob, "problem_image_png", None) else "problem",
            "pdf_lines": pdf_lines,
            "pdf_line_count": len(pdf_lines),
        }

    for prob in result.problems:
        image_paths: list[str] = []
        image_only_fallback = bool(prob.problem_image_png) and (
            not prob.text_reliable or not str(prob.text or "").strip()
        )
        if image_only_fallback and prob.problem_image_png:
            rp = _save_image_bytes(f"{stem_name}_q{prob.number}.png", prob.problem_image_png)
            if rp:
                image_paths.append(rp)
        for fig_index, fig_png in enumerate(prob.figure_pngs, start=1):
            rp = _save_image_bytes(f"{stem_name}_q{prob.number}_fig{fig_index}.png", fig_png)
            if rp:
                image_paths.append(rp)

        # 텍스트가 신뢰 불가인 문항만 image-only fallback으로 보존한다. 신뢰 가능한 PDF 라인은
        # stem/choices와 좌표 메타데이터를 유지해 이후 native equation 복원에 사용한다.
        if image_only_fallback:
            stem_text, choices = "", []
        else:
            stem_text, choices = _split_stem_and_choices(prob.text)
            pdf_line_geometries = list(getattr(prob, "line_geometries", []) or [])
            geometry_split = _split_stem_and_choices_from_pdf_geometry(
                prob.text,
                pdf_line_geometries,
            )
            if geometry_split is not None:
                geometry_stem, geometry_choices = geometry_split
                geometry_placeholder_count = _placeholder_count_in_fields(geometry_stem, geometry_choices)
                text_placeholder_count = _placeholder_count_in_fields(stem_text, choices)
                geometry_nonempty = sum(1 for choice in geometry_choices if str(choice or "").strip())
                text_nonempty = sum(1 for choice in choices if str(choice or "").strip())
                if (
                    (
                        len(geometry_choices) >= len(choices)
                        and geometry_placeholder_count < text_placeholder_count
                    )
                    or (
                        len(geometry_choices) > len(choices)
                        and geometry_placeholder_count <= text_placeholder_count
                        and geometry_nonempty >= max(text_nonempty, min(4, len(geometry_choices)))
                    )
                ):
                    stem_text, choices = geometry_stem, geometry_choices
            geometry_stem = _repair_pdf_stem_fractions_from_geometry(stem_text, pdf_line_geometries)
            if _placeholder_count_in_fields(geometry_stem, choices) < _placeholder_count_in_fields(stem_text, choices):
                stem_text = geometry_stem

        number = str(prob.number) if prob.number else ""
        loose_key = _recognized_pdf_loose_duplicate_key(number, stem_text) if loose_dedup_enabled else ""
        if loose_key and loose_key in loose_seen:
            sink.skipped += 1
            continue
        if loose_key:
            loose_seen.add(loose_key)

        title = f"{stem_name} #{number}" if number else stem_name
        sink.add(
            {
                **metadata,
                "source_type": "pdf",
                "source_name": filename,
                "source_page": prob.page_number,
                "number": number,
                "title": title,
                "stem": stem_text,
                "choices": choices,
                "image_paths": image_paths,
                "tables": [],
                "layout": layout_payload(prob),
            }
        )

    notices = [f"{len(sink.created)}개 문항을 PDF에서 자동 인식했습니다(문항 분리 + 그림/이미지 추출)."]
    if getattr(result, "exam_title", ""):
        notices.append(f"시험지 제목 감지: {result.exam_title}")
    notices.extend(result.notices)
    notices.extend(_dedup_notices(sink))
    return {"created": sink.created, "notices": notices, "exam_title": getattr(result, "exam_title", "")}


def _legacy_import_pdf(filename: str, payload: bytes, metadata: dict[str, Any]) -> dict[str, Any]:
    rel_path = save_upload(filename, payload)
    pdf_path = storage.DATA_DIR / rel_path
    sink = _Sink()
    notices: list[str] = []
    try:
        reader = PdfReader(str(pdf_path))
    except Exception as exc:  # pragma: no cover - library-specific errors
        problem = storage.create_problem(
            {
                **metadata,
                "source_type": "pdf",
                "source_name": filename,
                "title": Path(filename).stem,
                "stem": f"PDF를 열 수 없습니다: {exc}",
            }
        )
        return {"created": [problem], "notices": ["PDF 분석에 실패해 빈 문제로 등록했습니다."]}

    sequence = 1
    total_pdf_images = 0
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        page_images: list[str] = []
        try:
            for image_file in page.images:
                rel_path = _save_image_bytes(image_file.name or f"p{page_index}.png", image_file.data)
                if rel_path:
                    page_images.append(rel_path)
        except Exception:
            pass
        chunks = _split_questions(page_text)
        if not chunks:
            if page_images:
                sink.add(
                    {
                        **metadata,
                        "source_type": "pdf",
                        "source_name": filename,
                        "source_page": page_index,
                        "title": f"{Path(filename).stem} {page_index}쪽 이미지",
                        "stem": "",
                        "image_paths": page_images,
                    }
                )
                total_pdf_images += len(page_images)
            else:
                notices.append(f"{page_index}쪽에서 텍스트를 찾지 못했습니다.")
            continue
        for chunk_index, chunk in enumerate(chunks):
            number = _extract_number(chunk, sequence)
            sink.add(
                {
                    **metadata,
                    "source_type": "pdf",
                    "source_name": filename,
                    "source_page": page_index,
                    "number": number,
                    "title": f"{Path(filename).stem} #{number}",
                    "stem": chunk,
                    # 페이지 내 위치를 알 수 없어 페이지 첫 문항에 모아 붙인다.
                    "image_paths": page_images if chunk_index == 0 else [],
                }
            )
            sequence += 1
        total_pdf_images += len(page_images)
    if total_pdf_images:
        notices.append(f"PDF 이미지 {total_pdf_images}개를 페이지별 첫 문항에 첨부했습니다. 필요하면 편집에서 옮기세요.")
    notices.extend(_dedup_notices(sink))
    # 새로 만든 것도, 건너뛴 중복도 없을 때만 스캔 PDF 안내용 빈 문항을 만든다.
    if not sink.created and not sink.skipped:
        sink.add(
            {
                **metadata,
                "source_type": "pdf",
                "source_name": filename,
                "title": Path(filename).stem,
                "stem": "스캔 PDF이거나 텍스트를 추출하지 못했습니다. 이미지로 등록하거나 본문을 직접 입력하세요.",
            }
        )
    return {"created": sink.created, "notices": notices}


def import_image(filename: str, payload: bytes, metadata: dict[str, Any]) -> dict[str, Any]:
    rel_path = save_upload(filename, payload)
    full_path = storage.DATA_DIR / rel_path
    notices: list[str] = []
    width = height = 0
    try:
        with Image.open(full_path) as image:
            width, height = image.size
            image.verify()
    except Exception as exc:
        notices.append(f"이미지 확인 실패: {exc}")

    stem = metadata.get("stem") or ""
    has_real_text = bool(stem)
    if not stem and HAS_OCR:
        try:
            with Image.open(full_path) as image:
                ocr_text = pytesseract.image_to_string(image, lang="kor+eng")
            stem = _clean_text(ocr_text)
            if stem:
                has_real_text = True
                notices.append("OCR로 본문을 추출했습니다. 내용을 확인하세요.")
        except Exception as exc:
            notices.append(f"OCR 실패: {exc}")
    if not stem:
        # 식별 텍스트가 없으면 같은 안내 문구가 서로 중복으로 잡히지 않도록 중복 검사에서 뺀다.
        stem = "이미지 문항입니다. 본문이 필요하면 오른쪽 편집 영역에서 입력하세요."
    data = {
        **metadata,
        "source_type": "image",
        "source_name": filename,
        "title": metadata.get("title") or Path(filename).stem,
        "stem": stem,
        "image_paths": [rel_path],
    }
    problem = storage.create_problem_unique(data) if has_real_text else storage.create_problem(data)
    if problem is None:
        return {"created": [], "notices": ["이미 같은 본문의 문항이 있어 건너뛰었습니다.", *notices]}
    if width and height:
        notices.append(f"이미지 크기: {width}x{height}")
    return {"created": [problem], "notices": notices}


def import_text(
    filename: str,
    payload: bytes,
    metadata: dict[str, Any],
    source_type: str = "text",
) -> dict[str, Any]:
    save_upload(filename, payload)
    text = _decode_text(payload)
    chunks = _plain_text_chunks(text)
    if not chunks:
        return {"created": [], "notices": ["텍스트에서 문항을 찾지 못했습니다."]}
    sink = _create_from_chunks(chunks, source_type, filename, metadata)
    notices = [f"{len(sink.created)}개 문항을 텍스트에서 가져왔습니다."]
    notices.extend(_dedup_notices(sink))
    return {"created": sink.created, "notices": notices}


FIELD_ALIASES = {
    "number": ["number", "num", "no", "문항", "문항번호", "번호"],
    "subject": ["subject", "과목"],
    "unit": ["unit", "chapter", "단원", "소단원"],
    "tags": ["tags", "tag", "태그"],
    "title": ["title", "제목"],
    "stem": ["stem", "question", "body", "content", "본문", "문제", "문항본문"],
    "answer": ["answer", "정답"],
    "explanation": ["explanation", "solution", "해설", "풀이"],
}


def _first(row: dict[str, Any], key: str) -> str:
    aliases = FIELD_ALIASES.get(key, [key])
    lowered = {str(k).strip().lower(): v for k, v in row.items()}
    for alias in aliases:
        value = lowered.get(alias.lower())
        if value is not None:
            return str(value).strip()
    return ""


def _choices_from_row(row: dict[str, Any]) -> list[str]:
    choices: list[str] = []
    lowered = {str(k).strip().lower(): v for k, v in row.items()}
    for index in range(1, 10):
        for key in (f"choice{index}", f"option{index}", f"선지{index}", f"보기{index}"):
            if key.lower() in lowered and str(lowered[key.lower()]).strip():
                choices.append(str(lowered[key.lower()]).strip())
                break
    combined = lowered.get("choices") or lowered.get("options") or lowered.get("선지") or lowered.get("보기")
    if combined and not choices:
        choices = [
            part.strip()
            for part in re.split(r"\s*(?:\||;|\n)\s*", str(combined))
            if part.strip()
        ]
    return choices


def _split_inline_circled_choices(line: str) -> tuple[str, list[str]]:
    matches = list(INLINE_CIRCLED_CHOICE_RE.finditer(line))
    if len(matches) < 2:
        return line, []
    prefix = line[: matches[0].start()].strip()
    choices: list[str] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(line)
        value = line[start:end].strip()
        choices.append(value or match.group(1))
    if len(choices) < 2:
        return line, []
    return prefix, choices


def _split_inline_numeric_choices(line: str) -> tuple[str, list[str]]:
    matches = list(INLINE_NUMERIC_CHOICE_RE.finditer(line))
    if len(matches) < 3:
        return line, []
    labels = [int(match.group(1)) for match in matches]
    if labels[:3] != [1, 2, 3] or labels != sorted(labels) or len(set(labels)) != len(labels):
        return line, []
    prefix = line[: matches[0].start()].strip()
    choices: list[str] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(line)
        value = line[start:end].strip()
        choices.append(value or match.group(1))
    if len(choices) < 3:
        return line, []
    return prefix, choices


def _bare_numeric_choice_body(lines: list[str], index: int, choices_started: bool) -> str | None:
    match = BARE_NUMERIC_CHOICE_LINE_RE.match(lines[index].strip())
    if not match:
        return None
    label = int(match.group("label"))
    if label == 1 and not choices_started:
        lookahead: list[int] = []
        for raw in lines[index : index + 5]:
            next_match = BARE_NUMERIC_CHOICE_LINE_RE.match(raw.strip())
            if not next_match:
                break
            lookahead.append(int(next_match.group("label")))
        if lookahead[:3] != [1, 2, 3]:
            return None
    elif choices_started and label <= 5:
        pass
    else:
        return None
    return match.group("body").strip()


def _numeric_choice_label(label: str) -> int | None:
    match = NUMERIC_CHOICE_LABEL_RE.match(label.strip())
    return int(match.group(1)) if match else None


def _numeric_choice_sequence_starts(lines: list[str], index: int) -> bool:
    labels: list[int] = []
    for raw in lines[index : index + 5]:
        match = CHOICE_LINE_RE.match(raw.strip())
        if not match:
            break
        label = _numeric_choice_label(match.group("label"))
        if label is None:
            break
        labels.append(label)
    return labels[:3] == [1, 2, 3]


def _choice_line_body(
    lines: list[str],
    index: int,
    choices_started: bool,
    choice_count: int,
) -> str | None:
    line = lines[index].strip()
    match = CHOICE_LINE_RE.match(line)
    if not match:
        return None
    # 첫 줄의 "1. 문제..."는 문항 번호이므로 선지로 보지 않는다.
    if index == 0 and QUESTION_LINE_RE.match(line):
        return None
    label = match.group("label")
    number = _numeric_choice_label(label)
    if number is None:
        return match.group("body").strip()
    if not choices_started:
        if number != 1 or not _numeric_choice_sequence_starts(lines, index):
            return None
    elif number != choice_count + 1:
        return None
    return match.group("body").strip()


def _pdf_choice_fraction_part(value: str) -> bool:
    stripped = str(value or "").strip()
    if not stripped or len(stripped) > 80:
        return False
    if any(marker in stripped for marker in PDF_CHOICE_FRACTION_LABELS + PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS):
        return False
    if re.search(r"[\uac00-\ud7a3]", stripped):
        return False
    return bool(PDF_CHOICE_FRACTION_PART_RE.match(stripped))


def _pdf_choice_label_order(label: str) -> int:
    try:
        return PDF_CHOICE_FRACTION_LABELS.index(label)
    except ValueError:
        return 99


def _pdf_choice_exponent_base(value: str) -> bool:
    stripped = str(value or "").strip()
    return bool(stripped) and bool(re.search(r"[A-Za-z0-9α-ωΑ-Ωπ∞)\]}]", stripped))


def _pdf_choice_run(
    lines: list[str],
    index: int,
    pattern: re.Pattern[str],
) -> tuple[list[re.Match[str]], int] | None:
    matches: list[re.Match[str]] = []
    cursor = index
    while cursor < len(lines) and len(matches) < 5:
        match = pattern.match(lines[cursor].strip())
        if not match:
            break
        matches.append(match)
        cursor += 1
    labels = [match.group("label") for match in matches]
    if len(matches) < 2 or len(set(labels)) != len(labels):
        return None
    return matches, cursor


def _pdf_choice_denominators(
    lines: list[str],
    cursor: int,
    missing_count: int,
) -> tuple[list[str], int] | None:
    denominators: list[str] = []
    while (
        len(denominators) < missing_count
        and cursor < len(lines)
        and _pdf_choice_fraction_part(lines[cursor])
    ):
        denominators.append(lines[cursor].strip())
        cursor += 1
    if len(denominators) != missing_count:
        return None
    return denominators, cursor


def _render_pdf_choice_fraction_rows(
    matches: list[re.Match[str]],
    numerator_by_label: dict[str, str],
    denominator_by_label: dict[str, str],
    *,
    exponent: bool,
) -> list[str]:
    rows: list[str] = []
    for match in sorted(matches, key=lambda item: _pdf_choice_label_order(item.group("label"))):
        label = match.group("label")
        numerator = numerator_by_label[label]
        denominator = denominator_by_label[label]
        if exponent:
            rows.append(f"{label}{match.group('base')}^{{\\frac{{{numerator}}}{{{denominator}}}}}")
        else:
            sign = "-" if match.groupdict().get("sign") else ""
            rows.append(f"{label}{sign}\\frac{{{numerator}}}{{{denominator}}}")
    return rows


def _repair_pdf_choice_layout_after_numerators(
    lines: list[str],
    index: int,
    repaired: list[str],
    pattern: re.Pattern[str],
    *,
    exponent: bool,
) -> tuple[list[str], int, int] | None:
    run = _pdf_choice_run(lines, index, pattern)
    if run is None:
        return None
    matches, cursor = run
    count = len(matches)
    if exponent and not all(_pdf_choice_exponent_base(match.group("base")) for match in matches):
        return None
    if len(repaired) < count:
        return None
    numerator_rows = [line.strip() for line in repaired[-count:]]
    if not all(_pdf_choice_fraction_part(row) for row in numerator_rows):
        return None

    missing_count = sum(1 for match in matches if not (match.groupdict().get("den") or ""))
    denominator_rows, denominator_cursor = _pdf_choice_denominators(lines, cursor, missing_count) or (None, None)
    if denominator_rows is None or denominator_cursor is None:
        return None

    labels = sorted((match.group("label") for match in matches), key=_pdf_choice_label_order)
    numerator_by_label = {label: numerator_rows[offset] for offset, label in enumerate(labels)}
    denominator_iter = iter(denominator_rows)
    denominator_by_label: dict[str, str] = {}
    for match in matches:
        label = match.group("label")
        denominator_by_label[label] = (match.groupdict().get("den") or next(denominator_iter)).strip()

    return (
        _render_pdf_choice_fraction_rows(
            matches,
            numerator_by_label,
            denominator_by_label,
            exponent=exponent,
        ),
        denominator_cursor,
        count,
    )


def _repair_pdf_choice_layout_before_parts(
    lines: list[str],
    index: int,
    pattern: re.Pattern[str],
    *,
    exponent: bool,
) -> tuple[list[str], int, int] | None:
    run = _pdf_choice_run(lines, index, pattern)
    if run is None:
        return None
    matches, cursor = run
    count = len(matches)
    if exponent and not all(_pdf_choice_exponent_base(match.group("base")) for match in matches):
        return None
    if any(match.groupdict().get("den") for match in matches):
        return None
    numerator_rows = [line.strip() for line in lines[cursor : cursor + count]]
    if len(numerator_rows) != count or not all(_pdf_choice_fraction_part(row) for row in numerator_rows):
        return None
    denominator_cursor = cursor + count
    denominator_rows = [line.strip() for line in lines[denominator_cursor : denominator_cursor + count]]
    if len(denominator_rows) != count or not all(_pdf_choice_fraction_part(row) for row in denominator_rows):
        return None

    numerator_by_label = {
        match.group("label"): numerator_rows[offset] for offset, match in enumerate(matches)
    }
    denominator_by_label = {
        match.group("label"): denominator_rows[offset] for offset, match in enumerate(matches)
    }
    return (
        _render_pdf_choice_fraction_rows(
            matches,
            numerator_by_label,
            denominator_by_label,
            exponent=exponent,
        ),
        denominator_cursor + count,
        0,
    )


def _repair_pdf_choice_fraction_layout(text: str) -> str:
    """Rejoin PDF choice fractions split into numerator/placeholder/denominator rows."""
    lines = str(text or "").splitlines()
    if len(lines) < 6:
        return str(text or "")

    repaired: list[str] = []
    index = 0
    while index < len(lines):
        converted: tuple[list[str], int, int] | None = None
        for pattern, exponent in (
            (PDF_CHOICE_EXPONENT_FRACTION_RE, True),
            (PDF_CHOICE_FRACTION_RE, False),
        ):
            converted = _repair_pdf_choice_layout_after_numerators(
                lines,
                index,
                repaired,
                pattern,
                exponent=exponent,
            )
            if converted is not None:
                break
            converted = _repair_pdf_choice_layout_before_parts(
                lines,
                index,
                pattern,
                exponent=exponent,
            )
            if converted is not None:
                break

        if converted is not None:
            rows, next_index, pop_count = converted
            for _ in range(pop_count):
                if repaired:
                    repaired.pop()
            repaired.extend(rows)
            index = next_index
            continue

        repaired.append(lines[index])
        index += 1

    return "\n".join(repaired)


def _bbox_center(bbox: Any) -> tuple[float, float] | None:
    if not isinstance(bbox, list) or len(bbox) != 4:
        return None
    try:
        left, top, width, height = [float(value) for value in bbox]
    except (TypeError, ValueError):
        return None
    return left + width / 2.0, top + height / 2.0


def _bbox_top(bbox: Any) -> float | None:
    if not isinstance(bbox, list) or len(bbox) != 4:
        return None
    try:
        return float(bbox[1])
    except (TypeError, ValueError):
        return None


def _rect_center(rect: Any) -> tuple[float, float] | None:
    if not isinstance(rect, list) or len(rect) != 4:
        return None
    try:
        left, top, right, bottom = [float(value) for value in rect]
    except (TypeError, ValueError):
        return None
    return (left + right) / 2.0, (top + bottom) / 2.0


def _pdf_line_placeholder_centers(line: dict[str, Any]) -> list[tuple[float, float]]:
    centers: list[tuple[float, float]] = []
    for char in line.get("pdf_line_chars") or []:
        if not isinstance(char, dict):
            continue
        value = str(char.get("c") or "")
        normalized = normalize_recognized_math_text(value)
        if value not in PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS and not any(
            marker in normalized for marker in PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS
        ):
            continue
        center = _rect_center(char.get("bbox"))
        if center is not None:
            centers.append(center)
    return centers


def _pdf_stem_fraction_part(text: str) -> bool:
    stripped = str(text or "").strip()
    if not stripped or len(stripped) > 80:
        return False
    if any(marker in stripped for marker in PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS):
        return False
    if PDF_CHOICE_GEOMETRY_LABEL_RE.match(stripped):
        return False
    if re.search(r"[\uac00-\ud7a3]", stripped):
        return False
    if any(char in stripped for char in "<>=,"):
        return False
    if not re.search(r"[A-Za-z0-9α-ωΑ-Ωπ∞√∑∫]", stripped):
        return False
    if stripped in {"+", "-", "*", "/", "^", "_"}:
        return False
    return bool(PDF_STEM_FRACTION_PART_RE.match(stripped))


def _pdf_stem_fraction_head_allows(text: str, placeholder_offset: int) -> bool:
    head = str(text or "")[:placeholder_offset].rstrip()
    if not head:
        return True
    return head[-1] in "=+-*/([{,"


def _pdf_stem_fraction_suffix_allows(text: str, placeholder_offset: int) -> bool:
    suffix = str(text or "")[placeholder_offset + 1 :].strip()
    return suffix in {"", ")", "]", "}", ","}


def _matched_pdf_stem_entries(
    stem: str,
    line_geometries: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    source_lines = [line for line in str(stem or "").splitlines()]
    entries: list[dict[str, Any]] = []
    geometry_cursor = 0
    for stem_index, raw in enumerate(source_lines):
        text = raw.strip()
        if not text:
            continue
        while geometry_cursor < len(line_geometries):
            line = line_geometries[geometry_cursor]
            geometry_cursor += 1
            if not isinstance(line, dict):
                continue
            if str(line.get("text") or "").strip() != text:
                continue
            entries.append({"stem_index": stem_index, "geometry_index": geometry_cursor - 1, "line": line})
            break
    return entries


def _nearest_pdf_stem_fraction_part(
    entries: list[dict[str, Any]],
    *,
    stem_index: int,
    placeholder_center_x: float,
    placeholder_center_y: float,
    above: bool,
    used_stem_indices: set[int],
) -> dict[str, Any] | None:
    candidates: list[tuple[float, float, int, dict[str, Any]]] = []
    for entry in entries:
        candidate_stem_index = int(entry["stem_index"])
        if candidate_stem_index == stem_index or candidate_stem_index in used_stem_indices:
            continue
        line = entry["line"]
        text = str(line.get("text") or "").strip()
        if not _pdf_stem_fraction_part(text):
            continue
        center = _bbox_center(line.get("bbox_px"))
        if center is None:
            continue
        center_x, center_y = center
        x_distance = abs(center_x - placeholder_center_x)
        if x_distance > 70:
            continue
        vertical = placeholder_center_y - center_y if above else center_y - placeholder_center_y
        if vertical <= 0 or vertical > 100:
            continue
        candidates.append((vertical, x_distance, candidate_stem_index, entry))
    if not candidates:
        return None
    _, _, _, entry = sorted(candidates)[0]
    return entry


def _repair_pdf_stem_fractions_from_geometry(stem: str, line_geometries: list[dict[str, Any]]) -> str:
    """Recover simple stacked fractions in stems from exact placeholder char geometry."""
    if not stem or not line_geometries:
        return str(stem or "")
    entries = _matched_pdf_stem_entries(stem, line_geometries)
    if not entries:
        return str(stem or "")

    output_lines = [line for line in str(stem or "").splitlines()]
    used_stem_indices: set[int] = set()
    changed = False
    for entry in entries:
        stem_index = int(entry["stem_index"])
        if stem_index in used_stem_indices:
            continue
        line = entry["line"]
        text = str(line.get("text") or "").strip()
        placeholder_matches = list(re.finditer(r"[\u25a1\u25a2]", text))
        if not placeholder_matches:
            continue
        placeholder_centers = _pdf_line_placeholder_centers(line)
        if len(placeholder_centers) < len(placeholder_matches):
            continue

        replacements: list[tuple[int, int, str]] = []
        local_used: set[int] = set()
        for match, center in zip(placeholder_matches, placeholder_centers):
            if not _pdf_stem_fraction_head_allows(text, match.start()):
                continue
            if not _pdf_stem_fraction_suffix_allows(text, match.start()):
                continue
            numerator_entry = _nearest_pdf_stem_fraction_part(
                entries,
                stem_index=stem_index,
                placeholder_center_x=center[0],
                placeholder_center_y=center[1],
                above=True,
                used_stem_indices=used_stem_indices | local_used,
            )
            denominator_entry = _nearest_pdf_stem_fraction_part(
                entries,
                stem_index=stem_index,
                placeholder_center_x=center[0],
                placeholder_center_y=center[1],
                above=False,
                used_stem_indices=used_stem_indices | local_used,
            )
            if numerator_entry is None or denominator_entry is None:
                continue
            numerator = str(numerator_entry["line"].get("text") or "").strip()
            denominator = str(denominator_entry["line"].get("text") or "").strip()
            replacements.append((match.start(), match.end(), rf"\frac{{{numerator}}}{{{denominator}}}"))
            local_used.add(int(numerator_entry["stem_index"]))
            local_used.add(int(denominator_entry["stem_index"]))

        if not replacements:
            continue
        rebuilt = text
        for start, end, replacement in sorted(replacements, reverse=True):
            rebuilt = f"{rebuilt[:start]}{replacement}{rebuilt[end:]}"
        output_lines[stem_index] = rebuilt
        used_stem_indices.update(local_used)
        changed = True

    if not changed:
        return str(stem or "")
    return _clean_text("\n".join(line for index, line in enumerate(output_lines) if index not in used_stem_indices))


def _pdf_choice_geometry_part(text: str) -> bool:
    return _pdf_choice_fraction_part(text) and not PDF_CHOICE_GEOMETRY_LABEL_RE.match(str(text or "").strip())


def _pdf_choice_geometry_label_body(label_line: str) -> tuple[str, str] | None:
    match = PDF_CHOICE_GEOMETRY_LABEL_RE.match(str(label_line or "").strip())
    if not match:
        return None
    return match.group("label"), match.group("body").strip()


def _nearest_pdf_choice_part(
    indexed_lines: list[tuple[int, dict[str, Any]]],
    *,
    label_index: int,
    label_center_x: float,
    label_top: float,
    above: bool,
    used: set[int],
) -> tuple[int, str] | None:
    candidates: list[tuple[float, float, int, str]] = []
    for index, line in indexed_lines:
        if index == label_index or index in used:
            continue
        text = str(line.get("text") or "").strip()
        if not _pdf_choice_geometry_part(text):
            continue
        bbox = line.get("bbox_px")
        center = _bbox_center(bbox)
        top = _bbox_top(bbox)
        if center is None or top is None:
            continue
        center_x, _ = center
        x_distance = abs(center_x - label_center_x)
        if x_distance > 90:
            continue
        vertical = label_top - top if above else top - label_top
        if vertical <= 0 or vertical > 90:
            continue
        candidates.append((vertical, x_distance, index, text))
    if not candidates:
        return None
    _, _, index, text = sorted(candidates)[0]
    return index, text


def _pdf_choice_from_geometry_body(
    body: str,
    numerator: str | None,
    denominator: str | None,
) -> str:
    text = str(body or "").strip()
    sample_label = PDF_CHOICE_FRACTION_LABELS[0]

    def fraction(num: str, den: str, sign: str = "") -> str:
        return f"{sign}\\frac{{{num.strip()}}}{{{den.strip()}}}"

    if not any(char in text for char in PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS):
        body_denominator = text
        sign = ""
        if body_denominator.startswith(("-", "\u2212")):
            sign = "-"
            body_denominator = body_denominator[1:].strip()
        if numerator and body_denominator and _pdf_choice_fraction_part(body_denominator):
            return fraction(numerator, body_denominator, sign)
        if numerator and denominator and (not text or text in {"-", "\u2212", "+"}):
            return fraction(numerator, denominator, "-" if text in {"-", "\u2212"} else "")
        if numerator and denominator and not text:
            return fraction(numerator, denominator)
        return _normalize_pdf_choice_body(text)

    exponent_match = PDF_CHOICE_EXPONENT_FRACTION_RE.match(f"{sample_label}{text}")
    fraction_match = PDF_CHOICE_FRACTION_RE.match(f"{sample_label}{text}")
    if exponent_match and numerator and denominator and _pdf_choice_exponent_base(exponent_match.group("base")):
        return f"{exponent_match.group('base')}^{{\\frac{{{numerator}}}{{{denominator}}}}}"
    if fraction_match:
        sign = "-" if fraction_match.groupdict().get("sign") else ""
        inline_denominator = (fraction_match.groupdict().get("den") or "").strip()
        denominator = denominator or inline_denominator or None
        if numerator and denominator:
            return f"{sign}\\frac{{{numerator}}}{{{denominator}}}"
        if denominator:
            return f"{sign}\\frac{{1}}{{{denominator}}}"
    return text


def _split_stem_and_choices_from_pdf_geometry(
    text: str,
    line_geometries: list[dict[str, Any]],
) -> tuple[str, list[str]] | None:
    """Use PDF line bbox geometry to rejoin fraction choices split across rows."""
    if not line_geometries:
        return None
    indexed_lines = [
        (index, line)
        for index, line in enumerate(line_geometries)
        if isinstance(line, dict) and str(line.get("text") or "").strip()
    ]
    label_entries: list[dict[str, Any]] = []
    for index, line in indexed_lines:
        parsed = _pdf_choice_geometry_label_body(str(line.get("text") or ""))
        if parsed is None:
            continue
        label, body = parsed
        bbox = line.get("bbox_px")
        center = _bbox_center(bbox)
        top = _bbox_top(bbox)
        if center is None or top is None:
            continue
        label_entries.append(
            {
                "index": index,
                "label": label,
                "body": body,
                "center_x": center[0],
                "top": top,
            }
        )
    if len(label_entries) < 2:
        return None

    used: set[int] = {int(entry["index"]) for entry in label_entries}
    choices_by_label: dict[str, str] = {}
    for entry in sorted(label_entries, key=lambda item: _pdf_choice_label_order(str(item["label"]))):
        label = str(entry["label"])
        body = str(entry["body"])
        numerator: str | None = None
        denominator: str | None = None
        body_text = body.strip()
        has_placeholder = any(char in body_text for char in PDF_CHOICE_FRACTION_PLACEHOLDER_CHARS)
        body_is_sign = body_text in {"-", "\u2212", "+"}
        body_can_be_denominator = bool(
            body_text
            and not body_is_sign
            and _pdf_choice_fraction_part(body_text.lstrip("-\u2212+"))
        )
        needs_above_part = has_placeholder or not body_text or body_is_sign or body_can_be_denominator
        needs_below_part = has_placeholder or not body_text or body_is_sign
        if needs_above_part:
            numerator_item = _nearest_pdf_choice_part(
                indexed_lines,
                label_index=int(entry["index"]),
                label_center_x=float(entry["center_x"]),
                label_top=float(entry["top"]),
                above=True,
                used=used,
            )
            if numerator_item is not None:
                used.add(numerator_item[0])
                numerator = numerator_item[1]
        if needs_below_part:
            denominator_item = _nearest_pdf_choice_part(
                indexed_lines,
                label_index=int(entry["index"]),
                label_center_x=float(entry["center_x"]),
                label_top=float(entry["top"]),
                above=False,
                used=used,
            )
            if denominator_item is not None:
                used.add(denominator_item[0])
                denominator = denominator_item[1]
        choices_by_label[label] = _pdf_choice_from_geometry_body(body, numerator, denominator)

    labels = sorted(choices_by_label, key=_pdf_choice_label_order)
    if len(labels) < 2:
        return None
    choices = [choices_by_label[label] for label in labels]
    if not choices:
        return None

    stem_lines: list[str] = []
    for index, line in indexed_lines:
        if index in used:
            continue
        stem_lines.append(str(line.get("text") or "").strip())
    stem = re.sub(r"\n{3,}", "\n\n", "\n".join(stem_lines)).strip()
    return stem, [_normalize_pdf_choice_body(choice) for choice in choices]


def _placeholder_count_in_fields(stem: str, choices: list[str]) -> int:
    values = [stem, *choices]
    return sum(str(value or "").count("\u25a1") + str(value or "").count("\u25a2") for value in values)


def _normalize_pdf_choice_body(value: str) -> str:
    text = str(value or "").strip()
    match = PDF_CHOICE_UNIT_FRACTION_RE.match(text)
    if match:
        sign = "-" if match.group("sign") else ""
        return f"{sign}\\frac{{1}}{{{match.group('den')}}}"
    return text


def _split_stem_and_choices(text: str) -> tuple[str, list[str]]:
    """본문 안에 붙어 들어온 객관식 선지를 stem과 choices로 나눈다."""
    if not text:
        return "", []
    stem_lines: list[str] = []
    choices: list[str] = []
    lines = _repair_pdf_choice_fraction_layout(text).splitlines()
    for line_index, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line:
            stem_lines.append("")
            continue
        prefix, inline_choices = _split_inline_circled_choices(line)
        if inline_choices:
            if prefix:
                stem_lines.append(prefix)
            choices.extend(inline_choices)
            continue
        prefix, inline_choices = _split_inline_numeric_choices(line)
        if inline_choices:
            if prefix:
                stem_lines.append(prefix)
            choices.extend(inline_choices)
            continue
        choice_body = _choice_line_body(lines, line_index, bool(choices), len(choices))
        if choice_body is not None:
            choices.append(choice_body)
            continue
        bare_body = _bare_numeric_choice_body(lines, line_index, bool(choices))
        if bare_body is not None:
            choices.append(bare_body)
            continue
        stem_lines.append(raw_line.rstrip())
    return _clean_text("\n".join(stem_lines)), [_normalize_pdf_choice_body(choice) for choice in choices]


def _strip_leading_leaked_choice_block(stem: str, existing_choices: list[str]) -> str:
    """Drop a stale leading ①-⑤ block when HWP IR already has this problem's choices."""
    if not stem or len(existing_choices) < 2:
        return stem
    stripped = stem.lstrip()
    if not stripped or stripped[0] not in CIRCLED_CHOICE_MARKERS:
        return stem
    cleaned_stem, leaked_choices = _split_stem_and_choices(stem)
    if len(leaked_choices) < 5 or not cleaned_stem.strip():
        return stem
    return cleaned_stem


def _is_answer_heading(value: str) -> bool:
    normalized = re.sub(r"\s+", "", value)
    return normalized in {"및해설", "및풀이", "해설", "풀이"}


def _split_answer_explanation(text: str) -> tuple[str, str, str]:
    """본문 끝에 섞인 정답/해설 줄을 별도 필드로 분리한다."""
    body_lines: list[str] = []
    explanation_lines: list[str] = []
    answer = ""
    mode = "body"
    for raw_line in text.splitlines():
        line = raw_line.strip()
        answer_match = ANSWER_LINE_RE.match(line)
        if answer_match and not _is_answer_heading(answer_match.group("value")):
            answer = answer_match.group("value").strip()
            mode = "body"
            continue
        explanation_match = EXPLANATION_LINE_RE.match(line)
        if explanation_match:
            mode = "explanation"
            first = (explanation_match.group("value") or "").strip()
            if first:
                explanation_lines.append(first)
            continue
        if mode == "explanation":
            explanation_lines.append(raw_line.rstrip())
        else:
            body_lines.append(raw_line.rstrip())
    return (
        _clean_text("\n".join(body_lines)),
        answer,
        _clean_text("\n".join(explanation_lines)),
    )


def _problem_from_row(row: dict[str, Any], source_type: str, source_name: str) -> dict[str, Any] | None:
    stem = _first(row, "stem")
    title = _first(row, "title")
    if not stem and not title:
        return None
    choices = _choices_from_row(row)
    if stem and not choices:
        stem, extracted_answer, extracted_explanation = _split_answer_explanation(stem)
        stem, choices = _split_stem_and_choices(stem)
    else:
        extracted_answer = ""
        extracted_explanation = ""
    return {
        "source_type": source_type,
        "source_name": source_name,
        "number": _first(row, "number"),
        "subject": _first(row, "subject"),
        "unit": _first(row, "unit"),
        "tags": _first(row, "tags"),
        "title": title or (stem[:40] + ("..." if len(stem) > 40 else "")),
        "stem": stem,
        "choices": choices,
        "answer": _first(row, "answer") or extracted_answer,
        "explanation": _first(row, "explanation") or extracted_explanation,
    }


def import_csv(filename: str, payload: bytes, metadata: dict[str, Any]) -> dict[str, Any]:
    save_upload(filename, payload)
    text = _decode_text(payload)
    sample = text[:2048]
    dialect = csv.Sniffer().sniff(sample) if "," in sample or "\t" in sample else csv.excel
    reader = csv.DictReader(io.StringIO(text), dialect=dialect)
    sink = _Sink()
    for row in reader:
        problem_data = _problem_from_row(row, "csv", filename)
        if problem_data:
            sink.add({**metadata, **problem_data})
    return {
        "created": sink.created,
        "notices": [f"{len(sink.created)}개 문항을 가져왔습니다.", *_dedup_notices(sink)],
    }


def _sqlite_tables(conn: sqlite3.Connection) -> list[str]:
    rows = conn.execute(
        "SELECT name FROM sqlite_master WHERE type = 'table' AND name NOT LIKE 'sqlite_%'"
    ).fetchall()
    return [row[0] for row in rows]


def import_sqlite(filename: str, payload: bytes, metadata: dict[str, Any]) -> dict[str, Any]:
    rel_path = save_upload(filename, payload)
    src = storage.DATA_DIR / rel_path
    temp = storage.DATA_DIR / f"tmp_{uuid.uuid4().hex}.sqlite3"
    shutil.copyfile(src, temp)
    sink = _Sink()
    notices: list[str] = []
    try:
        conn = sqlite3.connect(temp)
        conn.row_factory = sqlite3.Row
        for table in _sqlite_tables(conn):
            columns = [row[1] for row in conn.execute(f"PRAGMA table_info({table})")]
            lower_columns = {column.lower() for column in columns}
            has_text = any(alias.lower() in lower_columns for alias in FIELD_ALIASES["stem"])
            has_title = any(alias.lower() in lower_columns for alias in FIELD_ALIASES["title"])
            if not has_text and not has_title:
                continue
            rows = conn.execute(f"SELECT * FROM {table} LIMIT 1000").fetchall()
            for row in rows:
                problem_data = _problem_from_row(dict(row), "sqlite", f"{filename}:{table}")
                if problem_data:
                    sink.add({**metadata, **problem_data})
            notices.append(f"{table}: {len(rows)}행 확인")
    finally:
        try:
            conn.close()  # type: ignore[name-defined]
        except Exception:
            pass
        temp.unlink(missing_ok=True)
    if not sink.created:
        notices.append("가져올 수 있는 question/stem/body/title 컬럼을 찾지 못했습니다.")
    notices.extend(_dedup_notices(sink))
    return {"created": sink.created, "notices": notices}


QUESTION_LINE_RE = re.compile(r"^\s*(?:문제\s*)?(\d{1,3})\s*[\.\)]")
PASSAGE_LEAD_RE = re.compile(r"^\s*\[\s*\d{1,2}\s*[~∼\-–]\s*\d{1,2}\s*\]")
QUESTION_TRAILER_RE = re.compile(
    r"^\s*\[(?P<score>\d+점)\]\[(?P<context>[^\]]*?(?P<number>\d{1,3}))\]\s*$"
)


def _save_image_bytes(name: str, payload: bytes) -> str | None:
    """이미지 바이트를 업로드 폴더에 저장하고 상대 경로를 돌려준다. 유효하지 않으면 None."""
    try:
        with Image.open(io.BytesIO(payload)) as image:
            image.verify()
    except Exception:
        return None
    return save_upload(name, payload)


def _first_child_math_text(element: Any, name: str) -> str:
    for child in element:
        if _local_name(child) == name:
            return _omml_text(child)
    return ""


def _first_descendant_math_text(element: Any, name: str) -> str:
    for child in element.iter():
        if child is element:
            continue
        if _local_name(child) == name:
            return _omml_text(child)
    return ""


def _child_math_texts(element: Any, name: str) -> list[str]:
    return [_omml_text(child) for child in element if _local_name(child) == name]


def _math_attr(element: Any, *names: str) -> str:
    wanted = set(names)
    for key, value in getattr(element, "attrib", {}).items():
        if str(key).rsplit("}", 1)[-1] in wanted and value:
            return str(value)
    return ""


def _first_descendant_math_attr(element: Any, child_name: str, *attr_names: str) -> str:
    for child in element.iter():
        if child is element:
            continue
        if _local_name(child) == child_name:
            value = _math_attr(child, *attr_names)
            if value:
                return value
    return ""


def _latex_group(value: str) -> str:
    text = value.strip()
    return text if text.startswith("{") and text.endswith("}") else "{" + text + "}"


def _latex_delimiter(value: str, default: str = ".") -> str:
    text = (value or default).strip() or default
    return {"{": r"\{", "}": r"\}", " ": "."}.get(text, text)


def _matrix_text(element: Any) -> str:
    rows: list[str] = []
    for row_node in element:
        if _local_name(row_node) != "mr":
            continue
        cells = [cell for cell in _child_math_texts(row_node, "e")]
        rows.append(" & ".join(cells))
    if not rows:
        return ""
    return r"\begin{matrix}" + r" \\ ".join(rows) + r"\end{matrix}"


def _omml_text(element: Any) -> str:
    """Word OMML 수식을 편집 가능한 LaTeX 스타일 텍스트 표현으로 바꾼다."""
    name = _local_name(element)
    if name == "t":
        return element.text or ""
    if name == "chr":
        return _math_attr(element, "val") or element.text or ""
    if name == "f":
        num = _first_child_math_text(element, "num")
        den = _first_child_math_text(element, "den")
        return f"\\frac{_latex_group(num)}{_latex_group(den)}" if num or den else ""
    if name == "rad":
        deg = _first_child_math_text(element, "deg")
        base = _first_child_math_text(element, "e")
        return f"\\sqrt[{deg}]{_latex_group(base)}" if deg else f"\\sqrt{_latex_group(base)}"
    if name == "d":
        body = _first_child_math_text(element, "e")
        beg = _latex_delimiter(_first_descendant_math_attr(element, "begChr", "val"), ".")
        end = _latex_delimiter(_first_descendant_math_attr(element, "endChr", "val"), ".")
        return f"\\left{beg}{body}\\right{end}" if body else ""
    if name == "bar":
        body = _first_child_math_text(element, "e")
        pos = _first_descendant_math_attr(element, "pos", "val")
        command = "underline" if pos == "bot" else "overline"
        return f"\\{command}{_latex_group(body)}" if body else ""
    if name == "sSup":
        base = _first_child_math_text(element, "e")
        sup = _first_child_math_text(element, "sup")
        return f"{base}^{_latex_group(sup)}" if sup else base
    if name == "sSub":
        base = _first_child_math_text(element, "e")
        sub = _first_child_math_text(element, "sub")
        return f"{base}_{_latex_group(sub)}" if sub else base
    if name == "sSubSup":
        base = _first_child_math_text(element, "e")
        sub = _first_child_math_text(element, "sub")
        sup = _first_child_math_text(element, "sup")
        if sub and sup:
            return f"{base}_{_latex_group(sub)}^{_latex_group(sup)}"
        return f"{base}_{_latex_group(sub)}" if sub else f"{base}^{_latex_group(sup)}"
    if name == "nary":
        symbol = _first_child_math_text(element, "chr") or _first_descendant_math_text(element, "chr") or "∑"
        sub = _first_child_math_text(element, "sub")
        sup = _first_child_math_text(element, "sup")
        body = _first_child_math_text(element, "e")
        symbol = {"∑": r"\sum", "Σ": r"\sum", "∫": r"\int", "∏": r"\prod"}.get(symbol, symbol)
        limit = ""
        if sub and sup:
            limit = f"_{_latex_group(sub)}^{_latex_group(sup)}"
        elif sub:
            limit = f"_{_latex_group(sub)}"
        elif sup:
            limit = f"^{_latex_group(sup)}"
        return f"{symbol}{limit} {body}".strip()
    if name == "acc":
        accent = _first_child_math_text(element, "chr") or _first_descendant_math_text(element, "chr")
        body = _first_child_math_text(element, "e")
        command = {
            "\u0305": "overline",
            "\u00af": "overline",
            "\u20d7": "vec",
            "\u2192": "vec",
            "\u0332": "underline",
        }.get(accent, "overline")
        return f"\\{command}{_latex_group(body)}" if body else ""
    if name == "limLow":
        base = _first_child_math_text(element, "e")
        limit = _first_child_math_text(element, "lim")
        command = r"\lim" if base.strip() == "lim" else base
        return f"{command}_{_latex_group(limit)}" if limit else command
    if name == "limUpp":
        base = _first_child_math_text(element, "e")
        limit = _first_child_math_text(element, "lim")
        command = r"\lim" if base.strip() == "lim" else base
        return f"{command}^{_latex_group(limit)}" if limit else command
    if name == "m":
        return _matrix_text(element)
    if name == "func":
        fname = _first_child_math_text(element, "fName")
        body = _first_child_math_text(element, "e")
        return f"{fname}({body})" if fname else body
    return "".join(_omml_text(child) for child in element)


def _docx_paragraph_text(para: Any) -> str:
    parts: list[str] = []
    for child in para._p.iterchildren():
        name = _local_name(child)
        if name in {"oMath", "oMathPara"}:
            math = _omml_text(child).strip()
            if math:
                parts.append(f"${math}$")
            continue
        if name == "r":
            run_text = "".join(
                node.text or ""
                for node in child.iter()
                if _local_name(node) == "t" and node.text
            )
            if run_text:
                parts.append(run_text)
    return "".join(parts).strip()


def _docx_cell_text(cell: Any) -> str:
    lines = [
        _docx_paragraph_text(paragraph).strip()
        for paragraph in getattr(cell, "paragraphs", []) or []
    ]
    return "\n".join(line for line in lines if line)


def _paragraphs_to_chunks(
    blocks: list[tuple[str, list[str], list[list[list[str]]]]],
) -> list[dict[str, Any]]:
    """(문단 텍스트, 이미지 경로들, 표들) 목록을 문항 번호 기준으로 묶는다."""
    chunks: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    pending_lead: dict[str, Any] | None = None

    def append_block(target: dict[str, Any], text: str, images: list[str], tables: list[list[list[str]]]) -> None:
        if text:
            target["text"].append(text)
        target["images"].extend(images)
        target["tables"].extend(tables)

    for text, images, tables in blocks:
        trailer_match = QUESTION_TRAILER_RE.match(text)
        if trailer_match and (current is not None or pending_lead is not None):
            if current is None:
                current = pending_lead or {"text": [], "images": [], "tables": []}
                pending_lead = None
                chunks.append(current)
            append_block(current, text, images, tables)
            current["number_hint"] = str(int(trailer_match.group("number")))
            current = None
            continue
        if QUESTION_LINE_RE.match(text):
            current = {
                "text": [],
                "images": [],
                "tables": [],
            }
            if pending_lead is not None:
                current["text"].extend(pending_lead["text"])
                current["images"].extend(pending_lead["images"])
                current["tables"].extend(pending_lead["tables"])
                pending_lead = None
            append_block(current, text, images, tables)
            chunks.append(current)
            continue
        if current is not None and PASSAGE_LEAD_RE.match(text):
            pending_lead = {"text": [], "images": [], "tables": []}
            append_block(pending_lead, text, images, tables)
            current = None
            continue
        if current is None:
            if pending_lead is None:
                pending_lead = {"text": [], "images": [], "tables": []}
            append_block(pending_lead, text, images, tables)
            continue
        append_block(current, text, images, tables)
    if pending_lead is not None and not chunks:
        chunks.append(pending_lead)
    result = []
    for chunk in chunks:
        body = _clean_text("\n".join(chunk["text"]))
        if body or chunk["images"] or chunk["tables"]:
            result.append(
                {
                    "text": body,
                    "images": chunk["images"],
                    "tables": chunk["tables"],
                    "number_hint": chunk.get("number_hint", ""),
                }
            )
    return result


def _create_from_chunks(
    chunks: list[dict[str, Any]],
    source_type: str,
    filename: str,
    metadata: dict[str, Any],
) -> _Sink:
    sink = _Sink()
    for sequence, chunk in enumerate(chunks, start=1):
        text, answer, explanation = _split_answer_explanation(chunk["text"])
        text, choices = _split_stem_and_choices(text)
        number = (
            str(chunk.get("number_hint") or "")
            or (_extract_number(text, sequence) if QUESTION_LINE_RE.match(text) else "")
        )
        title = f"{Path(filename).stem} #{number}" if number else Path(filename).stem
        sink.add(
            {
                **metadata,
                "source_type": source_type,
                "source_name": filename,
                "number": number,
                "title": title,
                "stem": text,
                "choices": choices,
                "answer": answer,
                "explanation": explanation,
                "image_paths": chunk["images"],
                "tables": chunk.get("tables") or [],
            }
        )
    return sink


VISUAL_CUE_RE = re.compile(r"(그림|그래프|도형|좌표평면|자료|표|곡선|직선|삼각형|사각형|원)")


def _attach_unpositioned_images(chunks: list[dict[str, Any]], images: list[str]) -> str:
    """위치 정보가 없는 HWP BinData 이미지를 문항들에 최대한 분산 배치한다."""
    if not chunks or not images:
        return ""
    if len(chunks) == 1:
        chunks[0]["images"] = [*images, *chunks[0]["images"]]
        return f"이미지 {len(images)}개는 위치를 알 수 없어 첫 문항에 첨부했습니다."

    cue_indexes = [
        index
        for index, chunk in enumerate(chunks)
        if VISUAL_CUE_RE.search(str(chunk.get("text") or ""))
    ]
    order: list[int] = []
    seen: set[int] = set()
    for index in cue_indexes:
        if index not in seen:
            order.append(index)
            seen.add(index)
    for index in range(len(chunks)):
        if index not in seen:
            order.append(index)
            seen.add(index)

    for image_index, image in enumerate(images):
        target = order[min(image_index, len(order) - 1)]
        chunks[target]["images"].append(image)
    return f"이미지 {len(images)}개는 HWP 내부 위치를 직접 읽지 못해 문항 순서 기준으로 분산 첨부했습니다."


def import_docx(filename: str, payload: bytes, metadata: dict[str, Any]) -> dict[str, Any]:
    from docx import Document as DocxDocument

    save_upload(filename, payload)
    notices: list[str] = []
    try:
        document = DocxDocument(io.BytesIO(payload))
    except Exception as exc:
        return {"created": [], "notices": [f"DOCX를 열 수 없습니다: {exc}"]}

    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    blip_ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    embed_attr = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    blocks: list[tuple[str, list[str], list[list[list[str]]]]] = []
    image_count = 0
    # 문단과 표를 문서 순서대로 훑어 표를 해당 문항에 정확히 붙인다.
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = Paragraph(child, document)
            images: list[str] = []
            for blip in para._p.iter(blip_ns):
                rel_id = blip.get(embed_attr)
                try:
                    part = (
                        document.part.rels[rel_id].target_part
                        if rel_id in document.part.rels
                        else None
                    )
                    if part is None:
                        continue
                    rel_path = _save_image_bytes(Path(part.partname).name, part.blob)
                except Exception:
                    continue
                if rel_path:
                    images.append(rel_path)
                    image_count += 1
            blocks.append((_docx_paragraph_text(para), images, []))
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            rows = [[_docx_cell_text(cell) for cell in row.cells] for row in table.rows]
            if any(any(cell for cell in row) for row in rows):
                blocks.append(("", [], [rows]))

    chunks = _paragraphs_to_chunks(blocks)
    if not chunks:
        return {"created": [], "notices": ["DOCX에서 내용을 찾지 못했습니다."]}
    sink = _create_from_chunks(chunks, "docx", filename, metadata)
    if image_count:
        notices.append(f"이미지 {image_count}개를 함께 가져왔습니다.")
    notices.extend(_dedup_notices(sink))
    return {"created": sink.created, "notices": notices}


def _local_name(element: Any) -> str:
    tag = element.tag
    return tag.rsplit("}", 1)[-1] if isinstance(tag, str) else ""


def _hwpx_table_rows(tbl: Any) -> list[list[str]]:
    """hp:tbl 엘리먼트에서 행×열 셀 텍스트를 뽑아 2차원 배열로 돌려준다."""
    rows: list[list[str]] = []
    for tr in tbl:
        if _local_name(tr) != "tr":
            continue
        cells: list[str] = []
        for tc in tr:
            if _local_name(tc) != "tc":
                continue
            texts: list[str] = []
            for node in tc.iter():
                name = _local_name(node)
                if name != "equation" and _inside_hwpx_container(node, {"equation"}):
                    continue
                if name == "t" and node.text:
                    texts.append(node.text)
                elif name == "equation":
                    texts.append(_hwpx_equation_text(node))
            cells.append("".join(texts).strip())
        if cells:
            rows.append(cells)
    return rows


def _inside_hwpx_container(node: Any, names: set[str]) -> bool:
    iter_ancestors = getattr(node, "iterancestors", None)
    if iter_ancestors is None:
        return False
    return any(_local_name(parent) in names for parent in iter_ancestors())


def _hwpx_equation_text(node: Any) -> str:
    for attr in ("script", "text", "formula", "eqn"):
        value = node.get(attr)
        if value:
            text = value.strip()
            return text if text.startswith("$") and text.endswith("$") else f"${text}$"
    for child in node.iter():
        if child is node or _local_name(child) != "script" or not child.text:
            continue
        text = str(child.text).strip()
        if text:
            return text if text.startswith("$") and text.endswith("$") else f"${text}$"
    texts = [
        str(child.text).strip()
        for child in node.iter()
        if child is not node and child.text and str(child.text).strip()
    ]
    text = " ".join(texts).strip()
    return f"${text}$" if text else "[수식 개체]"


def _hwpx_manifest_map(archive: zipfile.ZipFile) -> dict[str, str]:
    """content.hpf의 item id → zip 내 실제 경로 매핑."""
    mapping: dict[str, str] = {}
    names = set(archive.namelist())
    try:
        root = etree.fromstring(archive.read("Contents/content.hpf"))
    except Exception:
        return mapping
    for item in root.iter():
        if _local_name(item) != "item":
            continue
        item_id, href = item.get("id"), item.get("href")
        if not item_id or not href:
            continue
        candidates = [href, f"Contents/{href}", href.removeprefix("../"), f"BinData/{Path(href).name}"]
        for candidate in candidates:
            if candidate in names:
                mapping[item_id] = candidate
                break
    return mapping


def import_hwpx(filename: str, payload: bytes, metadata: dict[str, Any]) -> dict[str, Any]:
    save_upload(filename, payload)
    notices: list[str] = []
    try:
        archive = zipfile.ZipFile(io.BytesIO(payload))
    except Exception as exc:
        return {"created": [], "notices": [f"HWPX를 열 수 없습니다(zip 아님): {exc}"]}

    with archive:
        names = archive.namelist()
        sections = sorted(name for name in names if re.fullmatch(r"Contents/section\d+\.xml", name))
        if not sections:
            return {"created": [], "notices": ["HWPX 안에서 section XML을 찾지 못했습니다."]}
        bin_map = _hwpx_manifest_map(archive)
        saved_bins: dict[str, str | None] = {}
        paragraphs: list[tuple[str, list[str], list[list[list[str]]]]] = []
        image_count = 0
        for section_name in sections:
            try:
                root = etree.fromstring(archive.read(section_name))
            except Exception as exc:
                notices.append(f"{section_name} 분석 실패: {exc}")
                continue
            # 본문 문단(hp:p)만 직접 순회한다(셀 안의 중첩 문단은 표로 따로 처리).
            for para in root:
                if _local_name(para) != "p":
                    continue
                tables = [
                    _hwpx_table_rows(node) for node in para.iter() if _local_name(node) == "tbl"
                ]
                tables = [rows for rows in tables if rows]
                texts: list[str] = []
                images: list[str] = []
                for node in para.iter():
                    if _inside_hwpx_container(node, {"tbl", "equation"}):
                        continue
                    name = _local_name(node)
                    if name == "t" and node.text:
                        texts.append(node.text)
                    elif name == "equation":
                        texts.append(_hwpx_equation_text(node))
                    elif name == "img":
                        ref = node.get("binaryItemIDRef") or node.get("binaryItemIDRef".lower()) or ""
                        if ref not in saved_bins:
                            zip_path = bin_map.get(ref)
                            if zip_path is None:
                                # 매니페스트에 없으면 BinData에서 같은 이름을 추정
                                guess = [n for n in names if n.startswith("BinData/") and Path(n).stem == ref]
                                zip_path = guess[0] if guess else None
                            saved_bins[ref] = (
                                _save_image_bytes(Path(zip_path).name, archive.read(zip_path)) if zip_path else None
                            )
                        if saved_bins[ref]:
                            images.append(saved_bins[ref])
                            image_count += 1
                paragraphs.append(("".join(texts).strip(), images, tables))

    chunks = _paragraphs_to_chunks(paragraphs)
    if not chunks:
        return {"created": [], "notices": ["HWPX에서 내용을 찾지 못했습니다."]}
    sink = _create_from_chunks(chunks, "hwpx", filename, metadata)
    if image_count:
        notices.append(f"이미지 {image_count}개를 함께 가져왔습니다.")
    notices.extend(_dedup_notices(sink))
    return {"created": sink.created, "notices": notices}


# --- HWP(5.0 바이너리) 텍스트 추출 -------------------------------------------
# 문단 텍스트(HWPTAG_PARA_TEXT) 레코드의 UTF-16LE 문자열에서 컨트롤 문자를 걸러낸다.
# 코드 1~9, 11~12, 14~23은 8 WCHAR(16바이트)짜리 인라인/확장 컨트롤이다.
_HWP_EXTENDED = {1, 2, 3, 4, 5, 6, 7, 8, 11, 12, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23}
HWPTAG_PARA_TEXT = 67


def _hwp_decode_text(data: bytes) -> str:
    out: list[str] = []
    i = 0
    length = len(data) - 1
    while i < length:
        code = data[i] | (data[i + 1] << 8)
        if code == 9:
            out.append("\t")
            i += 16
        elif code in _HWP_EXTENDED:
            i += 16
        elif code in (10, 13):
            out.append("\n")
            i += 2
        elif code < 32:
            i += 2
        else:
            out.append(chr(code))
            i += 2
    return "".join(out)


def _hwp_iter_records(stream: bytes):
    pos = 0
    size_total = len(stream)
    while pos + 4 <= size_total:
        (header,) = struct.unpack_from("<I", stream, pos)
        tag = header & 0x3FF
        size = (header >> 20) & 0xFFF
        pos += 4
        if size == 0xFFF:
            if pos + 4 > size_total:
                break
            (size,) = struct.unpack_from("<I", stream, pos)
            pos += 4
        if pos + size > size_total:
            break
        yield tag, stream[pos : pos + size]
        pos += size


_ANSWER_SECTION_RE = re.compile(r"^\s*(?:\[정답\]|빠른\s*정답|정답\s*(?:및|과)\s*해설|정답\s*표)")


def _looks_like_answer_section(text: str) -> bool:
    """문항 뒤에 오는 정답표/해설 섹션인가(문항이 아니므로 제외 대상)."""
    t = (text or "").strip()
    if not t:
        return False
    return bool(_ANSWER_SECTION_RE.match(t)) or t.count("[정답]") >= 3


def _import_hwp_via_ir(
    filename: str, payload: bytes, metadata: dict[str, Any]
) -> dict[str, Any] | None:
    """rhwp to_ir() 로 원본 HWP → 편집 가능한 문항(텍스트+수식EQN+이미지). 실패 시 None."""
    try:
        from .importers_hwp_ir import hwp_to_problems
    except Exception:
        return None
    try:
        problems = hwp_to_problems(
            payload,
            filename,
            _save_image_bytes,
            _split_inline_circled_choices,
            chunk_paragraphs=_paragraphs_to_chunks,
            split_stem_choices=_split_stem_and_choices,
        )
    except Exception:
        return None
    if not problems:
        return None
    stem_name = Path(filename).stem
    sink = _Sink()
    answer_section = False
    for prob in problems:
        num = str(prob.get("number") or "")
        stem_text = prob.get("stem", "") or ""
        # 정답표·해설 섹션은 문항 뒤에 오므로, 한 번 만나면 이후 항목도 전부 제외한다.
        combined = stem_text + " " + " ".join(
            str(cell) for tbl in (prob.get("tables") or []) for row in tbl for cell in row
        )
        if answer_section or _looks_like_answer_section(combined):
            answer_section = True
            continue
        # 공유 지문은 원본 문제지처럼 테두리 박스로 낸다. writer 가 tables 를 테두리 표로
        # 렌더하므로, 안내문("[1~3] 다음 글을...")은 stem(박스 위), 지문 본문은 1칸 표(박스
        # 안)에 넣으면 편집 가능한 텍스트가 박스 안에 배치된다(writer 수정 불필요).
        if prob.get("is_passage"):
            head, _, body = stem_text.partition("\n")
            sink.add(
                {
                    **metadata,
                    "source_type": "hwp",
                    "source_name": filename,
                    "number": num,
                    "title": f"{stem_name} 지문{num}",
                    "unit": prob.get("unit") or metadata.get("unit", ""),
                    "stem": head.strip(),
                    "choices": [],
                    "image_paths": prob.get("image_paths", []),
                    "tables": [[[body.strip()]]] if body.strip() else [],
                }
            )
            continue
        # 번호가 비어 있으면 stem 선두("1." 등)에서 추출한다. 안 그러면 writer 가 순번
        # 인덱스를 붙여 "2. 1. …"처럼 중복 표기된다(지문이 순번 슬롯을 차지하므로).
        if not num and QUESTION_LINE_RE.match(stem_text):
            num = _extract_number(stem_text, "")
        choices = prob.get("choices", []) or []
        # 선지가 아직 stem 에 섞여 있으면(국어/영어 문단-청킹 경로) 분리해 번호·문제·보기를
        # 각각 편집 가능한 필드로 나눈다. 수학 마커 경로는 이미 선지가 분리돼 있어 건너뛴다.
        if not choices:
            stem_text, choices = _split_stem_and_choices(stem_text)
        else:
            stem_text = _strip_leading_leaked_choice_block(stem_text, choices)
        sink.add(
            {
                **metadata,
                "source_type": "hwp",
                "source_name": filename,
                "number": num,
                "title": f"{stem_name} #{num}" if num else stem_name,
                "unit": prob.get("unit") or metadata.get("unit", ""),
                "stem": stem_text,
                "choices": choices,
                "image_paths": prob.get("image_paths", []),
                "tables": prob.get("tables", []),
            }
        )
    notices = [f"{len(sink.created)}개 문항을 HWP에서 편집 가능하게 가져왔습니다(수식·이미지 포함)."]
    notices.extend(_dedup_notices(sink))
    return {"created": sink.created, "notices": notices}


def import_hwp(filename: str, payload: bytes, metadata: dict[str, Any]) -> dict[str, Any]:
    import olefile

    save_upload(filename, payload)
    # 1순위: rhwp to_ir() — 편집 가능한 [텍스트+수식(EQN)+이미지] 직접 추출. 수식은 $EQN$ 로
    #   stem 에 인라인 삽입돼 writer 의 native_math 가 hp:equation 으로 방출한다.
    #   실패/미설치면 아래 기존 경로(rhwp.paragraphs → OLE 레코드 → PrvText)로 폴백.
    ir_result = _import_hwp_via_ir(filename, payload, metadata)
    if ir_result is not None:
        return ir_result
    notices: list[str] = []
    try:
        ole = olefile.OleFileIO(io.BytesIO(payload))
    except Exception as exc:
        return {"created": [], "notices": [f"HWP 파일이 아닙니다: {exc}"]}

    with ole:
        if not ole.exists("FileHeader"):
            return {"created": [], "notices": ["HWP FileHeader가 없습니다. 한글 5.0 형식이 아닙니다."]}
        header = ole.openstream("FileHeader").read()
        flags = struct.unpack_from("<I", header, 36)[0] if len(header) >= 40 else 0
        compressed = bool(flags & 0x1)
        if flags & 0x2:
            return {"created": [], "notices": ["암호가 걸린 HWP는 가져올 수 없습니다."]}

        # 본문 텍스트 1순위: rhwp 엔진 (설치된 경우)
        paragraphs: list[tuple[str, list[str], list[list[list[str]]]]] = []
        if rhwp is not None:
            try:
                doc = rhwp.Document.from_bytes(payload)
                for para_text in doc.paragraphs():
                    for line in para_text.splitlines() or [""]:
                        paragraphs.append((line.strip(), [], []))
            except Exception:
                paragraphs = []

        # 2순위: BodyText/Section* 레코드 직접 파싱
        if not any(text for text, _, _ in paragraphs):
            paragraphs = []
            section_names = sorted(
                (entry for entry in ole.listdir() if len(entry) == 2 and entry[0] == "BodyText"),
                key=lambda entry: int(re.sub(r"\D", "", entry[1]) or 0),
            )
            for entry in section_names:
                raw = ole.openstream(entry).read()
                if compressed:
                    try:
                        raw = zlib.decompress(raw, -15)
                    except Exception as exc:
                        notices.append(f"{'/'.join(entry)} 압축 해제 실패: {exc}")
                        continue
                for tag, record in _hwp_iter_records(raw):
                    if tag != HWPTAG_PARA_TEXT:
                        continue
                    text = _hwp_decode_text(record)
                    for line in text.split("\n"):
                        paragraphs.append((line.strip(), [], []))

        # PrvText 보조: 본문 파싱이 실패했을 때 미리보기 텍스트라도 사용
        if not any(text for text, _, _ in paragraphs) and ole.exists("PrvText"):
            preview = ole.openstream("PrvText").read().decode("utf-16-le", errors="ignore")
            paragraphs = [(line.strip(), [], []) for line in preview.splitlines()]
            notices.append("본문 레코드 대신 미리보기 텍스트를 사용했습니다.")

        # 첨부 이미지: BinData 스토리지 전체 추출
        images: list[str] = []
        for entry in ole.listdir():
            if len(entry) != 2 or entry[0] != "BinData":
                continue
            blob = ole.openstream(entry).read()
            if compressed:
                try:
                    blob = zlib.decompress(blob, -15)
                except Exception:
                    pass
            rel_path = _save_image_bytes(entry[1], blob)
            if rel_path:
                images.append(rel_path)

    chunks = _paragraphs_to_chunks(paragraphs)
    if not chunks:
        return {"created": [], "notices": ["HWP에서 텍스트를 추출하지 못했습니다.", *notices]}
    if images:
        notice = _attach_unpositioned_images(chunks, images)
        if notice:
            notices.append(notice)
    sink = _create_from_chunks(chunks, "hwp", filename, metadata)
    notices.extend(_dedup_notices(sink))
    return {"created": sink.created, "notices": notices}
