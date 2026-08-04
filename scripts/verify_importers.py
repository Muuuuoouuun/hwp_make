# -*- coding: utf-8 -*-
"""임포터/내보내기 자체 검증 스크립트 (서버 없이 직접 호출)."""

from __future__ import annotations

import html
import io
import os
import re
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
os.environ["HWP_MAKE_DATA_DIR"] = str(ROOT / "data" / "verify_tmp2")

from PIL import Image  # noqa: E402

from app import docx_writer, exam_templates, hwpx_writer, hwpx_writer_v2, importers, importers_hwp_ir, math_text, storage  # noqa: E402
from app.recognition import pipeline as recognition_pipeline  # noqa: E402

# 중복 감지 검증이 결정적으로 동작하도록 매 실행마다 DB를 비운다.
storage.DB_PATH.unlink(missing_ok=True)
storage.init_db()

ROUNDTRIP_TEMPLATE_KEY = "school_exam"


def hwpx_text(path: Path) -> str:
    """내보낸 HWPX의 본문 텍스트를 zip/XML에서 직접 뽑는다(임포터·중복검사와 무관하게 출력 검증)."""
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("Contents/section0.xml").decode("utf-8")
    return html.unescape("".join(re.findall(r"<hp:t[^>]*>(.*?)</hp:t>", xml, re.S)))


def hwpx_section_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("Contents/section0.xml").decode("utf-8")


def hwpx_paragraph_texts(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("Contents/section0.xml"))
    hp_ns = "{http://www.hancom.co.kr/hwpml/2011/paragraph}"
    paragraphs: list[str] = []
    for para in root:
        if para.tag != f"{hp_ns}p":
            continue
        paragraphs.append("".join(node.text or "" for node in para.iter(f"{hp_ns}t")))
    return paragraphs


def hwpx_run_records(path: Path) -> list[dict[str, object]]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("Contents/section0.xml"))
    hp_ns = "{http://www.hancom.co.kr/hwpml/2011/paragraph}"
    records: list[dict[str, object]] = []

    def walk(node: ET.Element, in_table: bool = False) -> None:
        name = _xml_local_name(node.tag)
        child_in_table = in_table or name == "tbl"
        if name == "run":
            text = html.unescape("".join(child.text or "" for child in node.iter(f"{hp_ns}t")))
            records.append(
                {
                    "text": text,
                    "char": node.get("charPrIDRef") or "",
                    "in_table": in_table,
                }
            )
        for child in list(node):
            walk(child, child_in_table)

    walk(root)
    return records


def hwpx_has_distinct_math_run(path: Path, token: str, *, in_table: bool) -> bool:
    records = hwpx_run_records(path)
    matches = [
        record
        for record in records
        if record["in_table"] == in_table and token in str(record["text"]) and record["char"]
    ]
    plain_char_refs = {
        str(record["char"])
        for record in records
        if record["in_table"] == in_table
        and str(record["text"]).strip()
        and not math_text.extract_math_spans(str(record["text"]))
    }
    return any(str(record["char"]) not in plain_char_refs for record in matches)


def docx_text(path: Path) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def docx_document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _xml_local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if isinstance(tag, str) else ""


def docx_omml_texts(path: Path) -> list[str]:
    root = ET.fromstring(docx_document_xml(path))
    return [
        importers._omml_text(node)
        for node in root.iter()
        if _xml_local_name(node.tag) == "oMath"
    ]


def reset_problems() -> None:
    with storage.connect() as conn:
        conn.execute("DELETE FROM problems")
        conn.commit()


def make_png() -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (120, 60), (200, 30, 30)).save(buf, format="PNG")
    return buf.getvalue()


def make_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("1. 다음 중 광합성이 일어나는 장소는?")
    doc.add_paragraph("① 미토콘드리아  ② 엽록체  ③ 리보솜")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "세포소기관"
    table.rows[0].cells[1].text = "기능"
    table.rows[1].cells[0].text = "엽록체"
    table.rows[1].cells[1].text = "광합성"
    image_path = Path(tempfile.mkdtemp()) / "t.png"
    image_path.write_bytes(make_png())
    doc.add_picture(str(image_path))
    doc.add_paragraph("2. 세포 호흡의 결과 생성되는 물질은?")
    doc.add_paragraph("물과 이산화탄소에 대해 서술하시오.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_omml_table_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("41. 표 안의 수식을 확인하시오.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "항목"
    table.rows[0].cells[1].text = "값"
    table.rows[1].cells[0].text = "분수"
    math_node = docx_writer._omml_math(r"\frac{1}{2}")
    if math_node is not None:
        table.rows[1].cells[1].paragraphs[0]._p.append(math_node)
    doc.add_paragraph("정답: 1")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


RICH_FORMULAS = (
    "x^2",
    r"\frac{1}{2}",
    r"\frac{\sqrt{x}}{2}",
    "√",
    r"x_{n+1}",
    r"\sum_{k=1}^{n}",
    r"\sum_{k=1}^{n} k",
    r"\int_{0}^{1} dx",
    r"\overline{x}",
    r"\vec{v}",
    r"\left(x+1\right)",
    r"\begin{matrix}a & b \\ c & d\end{matrix}",
    r"\sqrt{x^2+1}",
    r"\sqrt[3]{x^2+1}",
    r"\binom{n}{r}",
    r"{}_{n}C_{r}",
    r"\left|x-1\right|",
    r"x\in\mathbb{R}",
    r"A\cup B",
    r"\overrightarrow{AB}",
    r"\widehat{ABC}",
    "α+β=γ",
    "≤",
    "≠",
)
RICH_LONG_PASSAGE = (
    "긴 지문은 인물이 낯선 도시의 골목을 지나며 오래된 약속을 떠올리는 장면으로 시작한다. "
    "화자는 빗소리와 가로등의 흔들림을 통해 기억이 현재의 선택을 밀어붙이는 과정을 차분히 서술한다. "
    "따라서 독자는 사건의 순서보다 정서의 변화와 서술자의 태도에 주목해야 한다."
)
RICH_STEM = "\n".join(
    [
        "3. 다음 수식 x^2, \\frac{1}{2}, \\frac{\\sqrt{x}}{2}, √, x_{n+1}, \\sum_{k=1}^{n}, \\sum_{k=1}^{n} k, \\int_{0}^{1} dx, \\overline{x}, \\vec{v}, \\left(x+1\\right), \\begin{matrix}a & b \\\\ c & d\\end{matrix}, \\sqrt{x^2+1}, \\sqrt[3]{x^2+1}, \\binom{n}{r}, {}_{n}C_{r}, \\left|x-1\\right|, x\\in\\mathbb{R}, A\\cup B, \\overrightarrow{AB}, \\widehat{ABC}, α+β=γ, a≤b, a≠0이 포함된 지문을 읽고 물음에 답하시오.",
        RICH_LONG_PASSAGE,
        "아래 표와 이미지를 참고하여 가장 적절한 설명을 고르시오.",
    ]
)
RICH_CHOICES = [
    r"x^2와 x_{n+1}는 시간의 흐름을 나타내는 핵심 단서이다.",
    r"\frac{1}{2}와 \sum_{k=1}^{n}은 화자의 양가적 태도를 상징한다.",
    r"√와 \sqrt{x^2+1}, a≤b, a≠0은 불확실한 기억이 현재로 이어지는 방식을 드러낸다.",
    "표와 이미지는 지문과 무관한 장식 자료이다.",
]
RICH_TABLES = [
    [
        ["요소", "유지 확인"],
        ["수식", "x^2 / \\frac{1}{2} / \\frac{\\sqrt{x}}{2} / √ / x_{n+1} / \\sum_{k=1}^{n} / \\sum_{k=1}^{n} k / \\int_{0}^{1} dx / \\overline{x} / \\vec{v} / \\left(x+1\\right) / \\begin{matrix}a & b \\\\ c & d\\end{matrix} / \\sqrt{x^2+1} / \\sqrt[3]{x^2+1} / \\binom{n}{r} / {}_{n}C_{r} / \\left|x-1\\right| / x\\in\\mathbb{R} / A\\cup B / \\overrightarrow{AB} / \\widehat{ABC} / α+β=γ / ≤ / ≠"],
        ["지문", "정서 변화"],
    ]
]
RICH_CHOICE_LABELS = ("①", "②", "③", "④")


failures = []

# 0a) Export template defaults used by the real UI/export path.
template_options = {template.key: template.export_option() for template in exam_templates.TEMPLATES}
native_math_default_templates = sorted(
    key for key, option in template_options.items() if option.get("native_math_default") is True
)
if native_math_default_templates != ["kice_math"]:
    failures.append(
        f"Export templates: native math should default on only for kice_math, got {native_math_default_templates}"
    )
if template_options.get("basic", {}).get("native_math_default") is not False:
    failures.append("Export templates: basic should keep native math off by default")

# 0) 공통 수식 인식/변환 유틸
math_spans = math_text.extract_math_spans(r"$x_{n+1}=\frac{1}{2}$, \frac{\sqrt{x}}{2}, \sum_{k=1}^{n} k, \int_{0}^{1} dx, \overline{x}, \vec{v}, \left(x+1\right), \begin{matrix}a & b \\ c & d\end{matrix}, \sin x+\cos x=1, \log_2 x, \alpha_1+\beta_2, \lim_{x\to0}\frac{\sin x}{x}, \sqrt[3]{x^2+1}, \binom{n}{r}, {}_{n}C_{r}, \left|x-1\right|, |x-1|, x\in\mathbb{R}, x\notin A, A\cup B, A\cap B, g\circ f, \overrightarrow{AB}, \widehat{ABC}, \mathrm{P}(A), f'(1)=3, x=-1, -2<x<3, 2x=4, 3(x-1)=6, f(x)=x^2+1, (x+1)/2=3, α²+βₙ, √x, √(x+1), √□5, sqrt□5, ∑_{k=1}^{n} k, ∫_0^1 f(x)dx, α+β=γ, a≤b, x_1=2, 한글x², 값α², 문장xₙ, 반지름r_1, x², xₙ, ½, ≤")
math_tokens = [span.text for span in math_spans]
math_normalized = {span.text: span.normalized for span in math_spans}
if (
    r"$x_{n+1}=\frac{1}{2}$" not in math_tokens
    or r"\frac{\sqrt{x}}{2}" not in math_tokens
    or r"\sum_{k=1}^{n} k" not in math_tokens
    or r"\int_{0}^{1} dx" not in math_tokens
    or r"\overline{x}" not in math_tokens
    or r"\vec{v}" not in math_tokens
    or r"\left(x+1\right)" not in math_tokens
    or r"\begin{matrix}a & b \\ c & d\end{matrix}" not in math_tokens
    or r"\sin x" not in math_tokens
    or r"\cos x" not in math_tokens
    or r"\log_2 x" not in math_tokens
    or r"\alpha_1" not in math_tokens
    or r"\beta_2" not in math_tokens
    or r"\lim_{x\to0}\frac{\sin x}{x}" not in math_tokens
    or r"\sqrt[3]{x^2+1}" not in math_tokens
    or r"\binom{n}{r}" not in math_tokens
    or r"{}_{n}C_{r}" not in math_tokens
    or r"\left|x-1\right|" not in math_tokens
    or r"|x-1|" not in math_tokens
    or r"x\in\mathbb{R}" not in math_tokens
    or r"x\notin A" not in math_tokens
    or r"A\cup B" not in math_tokens
    or r"A\cap B" not in math_tokens
    or r"g\circ f" not in math_tokens
    or r"\overrightarrow{AB}" not in math_tokens
    or r"\widehat{ABC}" not in math_tokens
    or r"\mathrm{P}(A)" not in math_tokens
    or "f'(1)=3" not in math_tokens
    or "x=-1" not in math_tokens
    or "-2<x<3" not in math_tokens
    or "2x=4" not in math_tokens
    or "3(x-1)=6" not in math_tokens
    or "f(x)=x^2+1" not in math_tokens
    or "(x+1)/2=3" not in math_tokens
    or "α²" not in math_tokens
    or "βₙ" not in math_tokens
    or "√x" not in math_tokens
    or "√(x+1)" not in math_tokens
    or "√□5" not in math_tokens
    or "sqrt□5" not in math_tokens
    or "∑_{k=1}^{n} k" not in math_tokens
    or "∫_0^1 f(x)dx" not in math_tokens
    or "α+β=γ" not in math_tokens
    or "a≤b" not in math_tokens
    or "x_1=2" not in math_tokens
    or "x²" not in math_tokens
    or "α²" not in math_tokens
    or "xₙ" not in math_tokens
    or "r_1" not in math_tokens
    or "x²" not in math_tokens
    or "xₙ" not in math_tokens
    or "½" not in math_tokens
    or "≤" not in math_tokens
):
    failures.append(f"Math analyzer: expected core tokens, got {math_tokens}")
for source, expected in {
    "한글x²": "x²",
    "값α²": "α²",
    "문장xₙ": "xₙ",
    "반지름r_1": "r_1",
}.items():
    if expected not in [span.text for span in math_text.extract_math_spans(source)]:
        failures.append(f"Math analyzer: Korean-adjacent formula {expected!r} missing in {source!r}")
if math_normalized.get("x²") != "x^{2}" or math_normalized.get("xₙ") != "x_{n}" or math_normalized.get("½") != r"\frac{1}{2}":
    failures.append(f"Math analyzer: unicode normalization wrong {math_normalized}")
if hwpx_writer._hancom_eqn_script("√x") != "sqrt {x}":
    failures.append(f"HWPX native math: unicode radical √x converted incorrectly: {hwpx_writer._hancom_eqn_script('√x')!r}")
if hwpx_writer._hancom_eqn_script("√(x+1)") != "sqrt {x+1}":
    failures.append(
        f"HWPX native math: unicode radical √(x+1) converted incorrectly: {hwpx_writer._hancom_eqn_script('√(x+1)')!r}"
    )
if math_text.split_math_text("√□5") != [("√□5", True)]:
    failures.append(f"Math analyzer: radical placeholder was split incorrectly {math_text.split_math_text('√□5')!r}")
if math_text.split_math_text("sqrt□5") != [("sqrt□5", True)]:
    failures.append(f"Math analyzer: Hancom radical placeholder was split incorrectly {math_text.split_math_text('sqrt□5')!r}")
if hwpx_writer._hancom_eqn_script("√□5") != "sqrt {5}" or hwpx_writer._hancom_eqn_script("sqrt□5") != "sqrt {5}":
    failures.append("HWPX native math: radical placeholder was not converted to a rooted argument")
if hwpx_writer._hancom_eqn_script("√□") != "sqrt {□}" or hwpx_writer._hancom_eqn_script("sqrt□") != "sqrt {□}":
    failures.append("HWPX native math: missing radical argument should stay as a square placeholder")
hwp_eqn_script = "f LEFT ( x RIGHT ) = {cases{``5x+a&&LEFT ( x<`-2 RIGHT )#``x ^{2} -a&&LEFT ( x GEQ `-2 RIGHT )}}"
if hwpx_writer._hancom_eqn_script(hwp_eqn_script) != hwp_eqn_script:
    failures.append("HWPX native math: source Hancom EQN script was altered instead of preserved")
hangul_cases_script = (
    "a _{n+1} = {cases{{a _{n}} over {3}&& LEFT ( a _{n} 이~3의~배수인~경우  RIGHT )"
    "#{a _{n}it^{2} +5} over {3}&& LEFT ( a _{n} 이~3의~배수가~아닌~경우  RIGHT )}}"
)
hangul_cases_text = importers_hwp_ir._wrap_eqn(hangul_cases_script)
if (
    "$a _{n+1} =$" not in hangul_cases_text
    or "${a _{n}} over {3}$ (a_n 이 3의 배수인 경우)" not in hangul_cases_text
    or "${a _{n}^{2} +5} over {3}$ (a_n 이 3의 배수가 아닌 경우)" not in hangul_cases_text
    or "cases{" in hangul_cases_text
):
    failures.append(f"HWP Hangul cases formula: condition split failed {hangul_cases_text!r}")
for segment, is_math in math_text.split_math_text(hangul_cases_text):
    if not is_math:
        continue
    converted_segment = hwpx_writer._hancom_eqn_script(segment) or ""
    if re.search(r"[\uac00-\ud7a3]", converted_segment):
        failures.append(f"HWP Hangul cases formula: Hangul leaked into equation segment {segment!r}")
    if "$" in converted_segment:
        failures.append(f"HWP Hangul cases formula: delimiter leaked into equation script {converted_segment!r}")
if math_text.extract_math_spans("시험일 2026-07-07"):
    failures.append("Math analyzer: date-like text was misread as formula")
for non_formula in ("쪽 범위 100-200", "문항 ID A1-2026", "가격은 $5 and $10"):
    if math_text.extract_math_spans(non_formula):
        failures.append(f"Math analyzer: non-formula text was misread as formula: {non_formula}")
# 통화 필터가 실제 수식(숫자+식별자)까지 삼키면 $ 짝이 밀려 수식/산문 스왑이 난다(2026-08-03).
for formula_like in ("주기가 $6  pi$이고", "반지름이 $60 r$일 때"):
    if not math_text.extract_math_spans(formula_like):
        failures.append(f"Math analyzer: currency filter swallowed a real formula: {formula_like}")

odd_form_key = importers._recognized_pdf_loose_duplicate_key(
    "26",
    "26. 그림과 같이 곡선 y=√□\n(미적분)\n홀수형\nx+xlnx와 x축 및 두 직선 x=1, x=2로 둘러싸인 부분을 밑면으로 하는 입체도형이 있다. [3점]\n① A",
)
even_form_key = importers._recognized_pdf_loose_duplicate_key(
    "26",
    "26. 그림과 같이 곡선 y=√□\n(미적분)\n짝수형\nx+xlnx와 x축 및 두 직선 x=1, x=2로 둘러싸인 부분을 밑면으로 하는 입체도형이 있다. [3점]\n① B",
)
if not odd_form_key or odd_form_key != even_form_key:
    failures.append(f"PDF import loose dedup: odd/even form keys differ {odd_form_key!r} vs {even_form_key!r}")
if odd_form_key == importers._recognized_pdf_loose_duplicate_key(
    "27",
    "27. 그림과 같이 곡선 y=√□ x+xlnx와 x축 및 두 직선 x=1, x=2로 둘러싸인 부분을 밑면으로 하는 입체도형이 있다. [3점]",
):
    failures.append("PDF import loose dedup: different problem numbers collapsed")

recognized_math = math_text.normalize_recognized_math_text(
    "\uf061+\uf062=\uf067, \U0001D465² ＋ \U0001D6FCₙ ＝ \U0001D7D9, \ue123\ue124"
)
if "α+β=γ" not in recognized_math or "x² + αₙ = 1" not in recognized_math or "□" not in recognized_math:
    failures.append(f"Recognized math normalization: unexpected output {recognized_math!r}")
recognized_vector = math_text.normalize_recognized_math_text("\ue06ea+\ue06eAB")
if recognized_vector != r"\vec{a}+\vec{AB}":
    failures.append(f"Recognized math normalization: vector accent PUA failed {recognized_vector!r}")
if hwpx_writer._hancom_eqn_script(recognized_vector) != "vec {a}+vec {AB}":
    failures.append(f"HWPX native math: vector accent PUA conversion failed {recognized_vector!r}")
recognized_wide_vector = math_text.normalize_recognized_math_text("\ue06d\ue06ea+\ue06d\ue06eAB")
if recognized_wide_vector != r"\vec{a}+\vec{AB}":
    failures.append(f"Recognized math normalization: wide vector accent PUA failed {recognized_wide_vector!r}")
recognized_trailing_vector = math_text.normalize_recognized_math_text("BC\ue046\ue06d\ue06e")
if recognized_trailing_vector != r"\vec{BC}":
    failures.append(
        f"Recognized math normalization: trailing vector accent PUA failed {recognized_trailing_vector!r}"
    )
if hwpx_writer._hancom_eqn_script(recognized_trailing_vector) != "vec {BC}":
    failures.append(
        f"HWPX native math: trailing vector accent PUA conversion failed {recognized_trailing_vector!r}"
    )
recognized_split_vector_residue = math_text.normalize_recognized_math_text(
    "|\u25a1\u20d7\n\\vec{XB}\nXC|"
)
if "\u25a1\u20d7" in recognized_split_vector_residue or not recognized_split_vector_residue.startswith(r"|\vec{XB}"):
    failures.append(
        "Recognized math normalization: split vector residue was not removed "
        f"{recognized_split_vector_residue!r}"
    )
recognized_radical = math_text.normalize_recognized_math_text("\ue05c\ue06d\ue036")
if recognized_radical != "√3":
    failures.append(f"Recognized math normalization: radical filler PUA failed {recognized_radical!r}")
if hwpx_writer._hancom_eqn_script(recognized_radical) != "sqrt {3}":
    failures.append(f"HWPX native math: radical filler PUA conversion failed {recognized_radical!r}")
recognized_segment = math_text.normalize_recognized_math_text("\ue06dAB=\ue06dBC")
if recognized_segment != r"\overline{AB}=\overline{BC}":
    failures.append(f"Recognized math normalization: segment overline PUA failed {recognized_segment!r}")
if hwpx_writer._hancom_eqn_script(recognized_segment) != "bar {AB}=bar {BC}":
    failures.append(f"HWPX native math: segment overline PUA conversion failed {recognized_segment!r}")
recognized_xbar = math_text.normalize_recognized_math_text("\ue06dx-c≤m≤\ue06dx+c")
if recognized_xbar != r"\overline{x}-c≤m≤\overline{x}+c":
    failures.append(f"Recognized math normalization: x-bar PUA failed {recognized_xbar!r}")
if hwpx_writer._hancom_eqn_script(recognized_xbar) != "bar {x}-c LEQ m LEQ bar {x}+c":
    failures.append(f"HWPX native math: x-bar PUA conversion failed {recognized_xbar!r}")
stacked_limit = math_text.normalize_recognized_math_layout_text(
    "lim\nh→0\ue06d\nf(1+h)-f(1)\uc758\nh\nvalue"
)
if "\\lim_{h\\to0}\\frac{f(1+h)-f(1)}{h}\uc758" not in stacked_limit:
    failures.append(f"Recognized math layout normalization: stacked limit fraction failed {stacked_limit!r}")
stacked_simple_fraction = math_text.normalize_recognized_math_layout_text("b=\ue06d2\n1\nok")
if "b=\\frac{1}{2}" not in stacked_simple_fraction:
    failures.append(f"Recognized math layout normalization: simple stacked fraction failed {stacked_simple_fraction!r}")
stacked_previous_fraction = math_text.normalize_recognized_math_layout_text("3t2-6t\nx=t3-\ue06d2\nok")
if "x=t3-\\frac{3t2-6t}{2}" not in stacked_previous_fraction:
    failures.append(f"Recognized math layout normalization: previous-line numerator fraction failed {stacked_previous_fraction!r}")
previous_text_guard = math_text.normalize_recognized_math_layout_text("\ubb38\uc7a5\nP(A)=\ue06d5")
if "\ubb38\uc7a5" not in previous_text_guard or r"P(A)=\frac{1}{5}" not in previous_text_guard:
    failures.append(f"Recognized math layout normalization: text previous line guard failed {previous_text_guard!r}")
previous_choice_guard = math_text.normalize_recognized_math_layout_text("\u2460\nx=\ue06d2")
if "\u2460\n" not in previous_choice_guard or r"x=\frac{1}{2}" not in previous_choice_guard:
    failures.append(f"Recognized math layout normalization: choice previous line guard failed {previous_choice_guard!r}")
unit_fraction_candidate = math_text.normalize_recognized_math_layout_text(
    "P(A)=\ue06d5\ncos(x)=-\ue06d3\nx\\to0\ue06d"
)
if r"P(A)=\frac{1}{5}" not in unit_fraction_candidate or r"cos(x)=-\frac{1}{3}" not in unit_fraction_candidate:
    failures.append(f"Recognized math layout normalization: unit numerator fractions failed {unit_fraction_candidate!r}")
if "x\\to0□" in unit_fraction_candidate or "x\\to0\ue06d" in unit_fraction_candidate:
    failures.append(f"Recognized math layout normalization: limit trailing placeholder cleanup failed {unit_fraction_candidate!r}")
choice_fraction_stem, choice_fraction_choices = importers._split_stem_and_choices(
    "\ubb38\uc81c\n7\n3\n4\n17\n9\n\u2460\ue06d\n\u2463\ue06d\n\u2464\ue06d\n\u2461\ue06d4\n\u2462\ue06d5\n10\n20\n10"
)
if choice_fraction_choices != [
    r"\frac{7}{10}",
    r"\frac{3}{4}",
    r"\frac{4}{5}",
    r"\frac{17}{20}",
    r"\frac{9}{10}",
]:
    failures.append(f"PDF choice fraction repair failed {choice_fraction_choices!r} stem={choice_fraction_stem!r}")
choice_fraction_guard_stem, choice_fraction_guard_choices = importers._split_stem_and_choices(
    "\ubb38\uc81c\n7\n3\n4\n17\n\u2460\ue06d\n\u2461\ue06d\n\u2462\ue06d\n\u2463\ue06d\n\u2464\ue06d\n10\n20\n10"
)
if len(choice_fraction_guard_choices) == 5 and any(r"\frac" in choice for choice in choice_fraction_guard_choices):
    failures.append(
        "PDF choice fraction repair over-converted an incomplete numerator block "
        f"{choice_fraction_guard_choices!r} stem={choice_fraction_guard_stem!r}"
    )
choice_exponent_source = math_text.normalize_recognized_math_layout_text(
    "문제\n13\n7\n5\n8\n17\n①2\ue06d\n⑤2\ue06d\n②2\ue06d3\n③2\ue06d2\n④2\ue06d3\n6\n6"
)
if "④2^{\\frac{6}{3}}" in choice_exponent_source:
    failures.append(f"Choice exponent guard: math_text pre-converted a grouped choice row {choice_exponent_source!r}")
choice_exponent_stem, choice_exponent_choices = importers._split_stem_and_choices(choice_exponent_source)
if choice_exponent_choices != [
    r"2^{\frac{13}{6}}",
    r"2^{\frac{7}{3}}",
    r"2^{\frac{5}{2}}",
    r"2^{\frac{8}{3}}",
    r"2^{\frac{17}{6}}",
]:
    failures.append(f"PDF choice exponent-fraction repair failed {choice_exponent_choices!r} stem={choice_exponent_stem!r}")
choice_forward_fraction_stem, choice_forward_fraction_choices = importers._split_stem_and_choices(
    "문제\n②\ue06d\n③\ue06d\n①\ue06d\n√3(3+8ln2)\n√3(5+12ln2)\n√3(1+12ln2)\n16\n24\n16"
)
if choice_forward_fraction_choices != [
    r"\frac{√3(1+12ln2)}{16}",
    r"\frac{√3(3+8ln2)}{16}",
    r"\frac{√3(5+12ln2)}{24}",
]:
    failures.append(
        f"PDF choice forward fraction repair failed {choice_forward_fraction_choices!r} "
        f"stem={choice_forward_fraction_stem!r}"
    )
choice_unit_fraction_source = math_text.normalize_recognized_math_layout_text(
    "문제\n①\ue06d5\n②-\ue06d3\n③2\ue06d3"
)
choice_unit_fraction_stem, choice_unit_fraction_choices = importers._split_stem_and_choices(choice_unit_fraction_source)
if choice_unit_fraction_choices != [r"\frac{1}{5}", r"-\frac{1}{3}", "2□3"]:
    failures.append(
        f"PDF choice unit fraction repair failed {choice_unit_fraction_choices!r} "
        f"stem={choice_unit_fraction_stem!r}"
    )


def _pdf_line(text: str, x: float, y: float, width: float = 48.0, height: float = 24.0) -> dict[str, object]:
    return {"text": text, "bbox_px": [x, y, width, height]}


def _pdf_placeholder_line(
    text: str,
    x: float,
    y: float,
    width: float,
    height: float,
    centers: list[tuple[float, float]],
) -> dict[str, object]:
    line = _pdf_line(text, x, y, width, height)
    line["pdf_line_chars"] = [
        {"c": "\u25a1", "bbox": [cx - 5.0, cy - 5.0, cx + 5.0, cy + 5.0]}
        for cx, cy in centers
    ]
    return line


choice_labels = ["\u2460", "\u2461", "\u2462", "\u2463", "\u2464"]
full_fraction_lines = [_pdf_line("15. stem", 0, 0, 500, 24)]
for offset, value in enumerate(("15", "27", "39", "51", "63")):
    full_fraction_lines.append(_pdf_line(value, 220 + offset * 120, 100))
for offset, label in enumerate(choice_labels):
    full_fraction_lines.append(_pdf_line(f"{label}\u25a1", 220 + offset * 120, 120, 66, 36))
for offset in range(5):
    full_fraction_lines.append(_pdf_line("4", 220 + offset * 120, 145))
full_fraction_split = importers._split_stem_and_choices_from_pdf_geometry(
    "\n".join(str(line["text"]) for line in full_fraction_lines),
    full_fraction_lines,
)
if full_fraction_split is None or full_fraction_split[1] != [
    r"\frac{15}{4}",
    r"\frac{27}{4}",
    r"\frac{39}{4}",
    r"\frac{51}{4}",
    r"\frac{63}{4}",
]:
    failures.append(f"PDF choice geometry repair: full fraction row failed {full_fraction_split!r}")

mixed_fraction_lines = [_pdf_line("6. stem", 0, 0, 500, 24)]
for x, value in ((220, "15"), (460, "17"), (700, "19")):
    mixed_fraction_lines.append(_pdf_line(value, x, 100))
for x, value in (
    (220, "\u2460\u25a1"),
    (340, "\u24618"),
    (460, "\u2462\u25a1"),
    (580, "\u24639"),
    (700, "\u2464\u25a1"),
):
    mixed_fraction_lines.append(_pdf_line(value, x, 120, 66, 36))
for x in (220, 460, 700):
    mixed_fraction_lines.append(_pdf_line("2", x, 145))
mixed_fraction_split = importers._split_stem_and_choices_from_pdf_geometry(
    "\n".join(str(line["text"]) for line in mixed_fraction_lines),
    mixed_fraction_lines,
)
if mixed_fraction_split is None or mixed_fraction_split[1] != [
    r"\frac{15}{2}",
    "8",
    r"\frac{17}{2}",
    "9",
    r"\frac{19}{2}",
]:
    failures.append(f"PDF choice geometry repair: mixed fraction row failed {mixed_fraction_split!r}")

stem_curve_lines = [
    _pdf_placeholder_line("26. curve y=\u25a1", 0, 100, 220, 80, [(108, 130)]),
    _pdf_line("3", 100, 95, 16, 20),
    _pdf_line("(x>1) meets", 140, 115, 120, 20),
    _pdf_line("x-1", 90, 145, 36, 20),
    _pdf_line("3", 300, 205, 16, 20),
    _pdf_placeholder_line("again y=\u25a1", 200, 200, 220, 60, [(308, 230)]),
    _pdf_line("(x>1)", 340, 215, 60, 20),
    _pdf_line("x-1", 290, 245, 36, 20),
]
stem_curve_repaired = importers._repair_pdf_stem_fractions_from_geometry(
    "\n".join(str(line["text"]) for line in stem_curve_lines),
    stem_curve_lines,
)
if stem_curve_repaired != "\n".join(
    [
        r"26. curve y=\frac{3}{x-1}",
        "(x>1) meets",
        r"again y=\frac{3}{x-1}",
        "(x>1)",
    ]
):
    failures.append(f"PDF stem geometry repair: curve fraction failed {stem_curve_repaired!r}")

stem_hyperbola_lines = [
    _pdf_line("x2", 100, 100, 20, 20),
    _pdf_line("y2", 160, 100, 20, 20),
    _pdf_placeholder_line("\u25a1", 95, 120, 30, 30, [(110, 135)]),
    _pdf_placeholder_line("-\u25a1", 145, 120, 45, 30, [(170, 135)]),
    _pdf_line("=-1", 200, 125, 50, 20),
    _pdf_line("9", 105, 150, 12, 20),
    _pdf_line("16", 160, 150, 24, 20),
]
stem_hyperbola_repaired = importers._repair_pdf_stem_fractions_from_geometry(
    "\n".join(str(line["text"]) for line in stem_hyperbola_lines),
    stem_hyperbola_lines,
)
if stem_hyperbola_repaired != "\n".join([r"\frac{x2}{9}", r"-\frac{y2}{16}", "=-1"]):
    failures.append(f"PDF stem geometry repair: hyperbola fraction failed {stem_hyperbola_repaired!r}")


def _pdf_char(c: str, left: float, top: float, right: float, bottom: float) -> dict[str, object]:
    return {"c": c, "bbox": [left, top, right, bottom]}


rich_accent_fraction_lines = [
    {
        **_pdf_line("14. □", 0, 0, 80, 35),
        "pdf_line_chars": [_pdf_char("\ue06d", 20, 10, 50, 30)],
    },
    {
        **_pdf_line("AB=3, □", 20, 10, 130, 35),
        "pdf_line_chars": [
            _pdf_char("A", 20, 23, 34, 39),
            _pdf_char("B", 34, 23, 48, 39),
            _pdf_char("\ue06d", 100, 10, 128, 30),
        ],
    },
    {
        **_pdf_line("π triangle", 200, 0, 100, 25),
        "pdf_line_chars": [_pdf_char("π", 204, 10, 218, 24)],
    },
    {
        **_pdf_line(r"BC=4, angle=\frac{1}{2}", 100, 20, 150, 35),
        "pdf_line_chars": [
            _pdf_char("B", 100, 23, 114, 39),
            _pdf_char("C", 114, 23, 128, 39),
            _pdf_char("\ue06d", 200, 25, 224, 45),
            _pdf_char("2", 205, 36, 217, 50),
        ],
    },
]
rich_accent_fraction_repaired = importers._repair_pdf_stem_fractions_from_geometry(
    "\n".join(str(line["text"]) for line in rich_accent_fraction_lines),
    rich_accent_fraction_lines,
)
if (
    r"\overline{AB}" not in rich_accent_fraction_repaired
    or r"\overline{BC}" not in rich_accent_fraction_repaired
    or r"\frac{π}{2}" not in rich_accent_fraction_repaired
    or "ABC" in rich_accent_fraction_repaired
    or "□" in rich_accent_fraction_repaired
):
    failures.append(
        "PDF stem character geometry repair: accent/fraction pairing failed "
        f"{rich_accent_fraction_repaired!r}"
    )

rich_two_fraction_lines = [
    {
        **_pdf_line("a4+□", 100, 100, 100, 35),
        "pdf_line_chars": [
            _pdf_char("a", 100, 92, 112, 108),
            _pdf_char("4", 112, 96, 122, 108),
            _pdf_char("\ue06d", 160, 100, 188, 124),
        ],
    },
    {
        **_pdf_line("a2=30", 160, 82, 80, 25),
        "pdf_line_chars": [_pdf_char("a", 162, 83, 174, 99), _pdf_char("2", 174, 87, 184, 99)],
    },
    {
        **_pdf_line("□", 95, 100, 30, 24),
        "pdf_line_chars": [_pdf_char("\ue06d", 96, 100, 124, 124)],
    },
    {
        **_pdf_line("a2", 95, 116, 30, 24),
        "pdf_line_chars": [_pdf_char("a", 98, 116, 110, 132), _pdf_char("2", 110, 120, 120, 132)],
    },
    {
        **_pdf_line("a1", 160, 116, 30, 24),
        "pdf_line_chars": [_pdf_char("a", 162, 116, 174, 132), _pdf_char("1", 174, 120, 184, 132)],
    },
]
rich_two_fraction_repaired = importers._repair_pdf_stem_fractions_from_geometry(
    "\n".join(str(line["text"]) for line in rich_two_fraction_lines),
    rich_two_fraction_lines,
)
if (
    r"\frac{a4}{a2}" not in rich_two_fraction_repaired
    or r"\frac{a2}{a1}" not in rich_two_fraction_repaired
    or "□" in rich_two_fraction_repaired
):
    failures.append(
        "PDF stem character geometry repair: adjacent fractions crossed "
        f"{rich_two_fraction_repaired!r}"
    )

intentional_blank_lines = [
    {
        **_pdf_line("(가) □", 0, 0, 120, 35),
        "pdf_line_chars": [
            _pdf_char("가", 0, 10, 18, 30),
            _pdf_char("\ue06d", 60, 5, 105, 35),
        ],
    }
]
intentional_blank_repaired = importers._repair_pdf_stem_fractions_from_geometry(
    "(가) □",
    intentional_blank_lines,
)
if intentional_blank_repaired != "(가) □":
    failures.append(
        "PDF stem character geometry repair: intentional worksheet blank changed "
        f"{intentional_blank_repaired!r}"
    )
p_over_q_placeholder = math_text.normalize_recognized_math_layout_text(
    "\ud655\ub960\uc740 \ue06dp\np+q\uc758 \uac12\uc744 \uad6c\ud558\uc2dc\uc624. (\ub2e8, p\uc640 q\ub294 \uc11c\ub85c\uc18c)"
)
if "\\frac{p}{q}" not in p_over_q_placeholder:
    failures.append(f"Recognized math layout normalization: p/q placeholder failed {p_over_q_placeholder!r}")
p_over_q_context_guard = math_text.normalize_recognized_math_layout_text("\ue06dp\np+q")
if "\\frac" in p_over_q_context_guard:
    failures.append(f"Recognized math layout normalization: p/q placeholder missing context guard failed {p_over_q_context_guard!r}")
p_over_q_radical_guard = math_text.normalize_recognized_math_layout_text(
    "q√5\uc774\ub2e4. p+q\uc758 \uac12\uc744 \uad6c\ud558\uc2dc\uc624.\n\ub113\uc774\ub294 \ue06dp\n(\ub2e8, p\uc640 q\ub294 \uc11c\ub85c\uc18c)"
)
if "\\frac{p}{q}" in p_over_q_radical_guard:
    failures.append(f"Recognized math layout normalization: p/q radical context was over-converted {p_over_q_radical_guard!r}")
log_base_placeholder = math_text.normalize_recognized_math_layout_text("log3\ue06da+log10\ue06dx")
if r"\log_{3}a+\log_{10}x" not in log_base_placeholder:
    failures.append(f"Recognized math layout normalization: log base placeholder failed {log_base_placeholder!r}")
log_base_word_guard = math_text.normalize_recognized_math_layout_text("blog3\ue06da")
if r"\log_{3}a" in log_base_word_guard:
    failures.append(f"Recognized math layout normalization: log base placeholder over-converted word text {log_base_word_guard!r}")
stacked_exponent_candidate = math_text.normalize_recognized_math_layout_text("9\ue06d4\n1")
if "9^{\\frac{1}{4}}" not in stacked_exponent_candidate:
    failures.append(f"Recognized math layout normalization: stacked fraction exponent failed {stacked_exponent_candidate!r}")
stacked_previous_exponent = math_text.normalize_recognized_math_layout_text("1\n×2\ue06d2")
if "×2^{\\frac{1}{2}}" not in stacked_previous_exponent:
    failures.append(f"Recognized math layout normalization: previous-line stacked fraction exponent failed {stacked_previous_exponent!r}")
stacked_delayed_base_exponent = math_text.normalize_recognized_math_layout_text("-\\frac{1}{2}\n× 3")
if "× 3^{-\\frac{1}{2}}" not in stacked_delayed_base_exponent:
    failures.append(f"Recognized math layout normalization: delayed exponent base repair failed {stacked_delayed_base_exponent!r}")
split_nth_root_exponent = math_text.normalize_recognized_math_layout_text("1. \u221a\n1\n3\ue06d5\u00d7 25\ue06d3")
if "\\sqrt[3]{5}\u00d7 25^{\\frac{1}{3}}" not in split_nth_root_exponent:
    failures.append(
        f"Recognized math layout normalization: split nth-root exponent failed {split_nth_root_exponent!r}"
    )
stacked_exponent_text_guard = math_text.normalize_recognized_math_layout_text("\uc9c1\uc120 y=3\ue06d13\n2\ub85c")
if "^{\\frac{2}{13}}" in stacked_exponent_text_guard:
    failures.append(f"Recognized math layout normalization: text line was over-converted to exponent {stacked_exponent_text_guard!r}")
stacked_root_candidate = math_text.normalize_recognized_math_layout_text("√\n\ue06d\nx+1")
if r"\overline{x}" in stacked_root_candidate:
    failures.append(f"Recognized math layout normalization: split radical was over-converted to x-bar {stacked_root_candidate!r}")
case_tail = math_text.normalize_recognized_math_layout_text("f(x)=\n{\nx\n(x>0)\n\ue07a\n이다")
if "□" in case_tail:
    failures.append(f"Recognized math layout normalization: cases tail placeholder was not removed {case_tail!r}")
if recognition_pipeline._pua_ratio("\uf061+\uf062=\uf067") != 0:
    failures.append("Recognition PUA ratio: recoverable Symbol-font math was counted as broken")
if recognition_pipeline._pua_ratio("\ue06ea+\ue06eAB") != 0:
    failures.append("Recognition PUA ratio: recoverable Hancom vector accent was counted as broken")
if recognition_pipeline._pua_ratio("\ue123x") <= 0:
    failures.append("Recognition PUA ratio: unknown PUA was not counted as broken")
if recognition_pipeline._clean_pua("\uf061\ue123\ue124\uf062") != "α□β":
    failures.append("Recognition PUA clean: unknown PUA run was not collapsed between recovered symbols")

stored_math = storage.create_problem(
    {
        "source_type": "manual",
        "source_name": "recognized_math",
        "title": "정규화",
        "stem": "1. \U0001D465² ＋ \uf061ₙ ＝ \U0001D7D9",
        "choices": ["\uf062≤\uf067"],
        "answer": "\uf061",
        "tables": [[["값", "\U0001D6FCₙ"]]],
    }
)
if "x² + αₙ = 1" not in stored_math["stem"] or stored_math["choices"] != ["β≤γ"] or stored_math["answer"] != "α":
    failures.append(f"Storage normalization: recognized math was not normalized on save {stored_math!r}")
if stored_math.get("math_summary", {}).get("count", 0) < 3:
    failures.append(f"Storage normalization: normalized math was not analyzed {stored_math.get('math_summary')}")
updated_math = storage.update_problem(
    stored_math["id"],
    {
        "stem": "2. \U0001D465 ＋ \uf061 ＝ \U0001D7DA",
        "choices": ["\uf061", "\uf062"],
        "tables": [[["수식", "\uf067²"]]],
    },
)
if "x + α = 2" not in updated_math["stem"] or updated_math["choices"] != ["α", "β"] or "γ²" not in str(updated_math["tables"]):
    failures.append(f"Storage normalization: recognized math was not normalized on update {updated_math!r}")
reset_problems()

omml = ET.fromstring(
    """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:f><m:num><m:r><m:t>1</m:t></m:r></m:num><m:den><m:r><m:t>2</m:t></m:r></m:den></m:f>
      <m:sSubSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sub><m:r><m:t>n</m:t></m:r></m:sub><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSubSup>
      <m:nary><m:chr m:val="∑"/><m:sub><m:r><m:t>k=1</m:t></m:r></m:sub><m:sup><m:r><m:t>n</m:t></m:r></m:sup><m:e><m:r><m:t>k</m:t></m:r></m:e></m:nary>
    </m:oMath>"""
)
omml_text = importers._omml_text(omml)
if r"\frac{1}{2}" not in omml_text or "x_{n}^{2}" not in omml_text or r"\sum_{k=1}^{n}" not in omml_text:
    failures.append(f"OMML conversion: LaTeX-like output missing in {omml_text!r}")

omml_narypr = ET.fromstring(
    """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:nary><m:naryPr><m:chr m:val="∫"/></m:naryPr><m:sub><m:r><m:t>0</m:t></m:r></m:sub><m:sup><m:r><m:t>1</m:t></m:r></m:sup><m:e><m:r><m:t>dx</m:t></m:r></m:e></m:nary>
    </m:oMath>"""
)
if r"\int_{0}^{1}" not in importers._omml_text(omml_narypr):
    failures.append("OMML conversion: naryPr/chr integral was not converted")

omml_accent = ET.fromstring(
    """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:acc><m:accPr><m:chr m:val="&#x20d7;"/></m:accPr><m:e><m:r><m:t>v</m:t></m:r></m:e></m:acc>
    </m:oMath>"""
)
if r"\vec{v}" not in importers._omml_text(omml_accent):
    failures.append("OMML conversion: vector accent was not converted")

omml_delimiter = ET.fromstring(
    """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:d><m:dPr><m:begChr m:val="("/><m:endChr m:val=")"/></m:dPr><m:e><m:r><m:t>x+1</m:t></m:r></m:e></m:d>
    </m:oMath>"""
)
if r"\left(x+1\right)" not in importers._omml_text(omml_delimiter):
    failures.append("OMML conversion: delimiter group was not converted")

omml_bar = ET.fromstring(
    """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:bar><m:barPr><m:pos m:val="top"/></m:barPr><m:e><m:r><m:t>AB</m:t></m:r></m:e></m:bar>
    </m:oMath>"""
)
if r"\overline{AB}" not in importers._omml_text(omml_bar):
    failures.append("OMML conversion: bar/overline was not converted")

omml_limit = ET.fromstring(
    """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:limLow><m:e><m:r><m:t>lim</m:t></m:r></m:e><m:lim><m:r><m:t>x→0</m:t></m:r></m:lim></m:limLow>
      <m:f><m:num><m:r><m:t>sin x</m:t></m:r></m:num><m:den><m:r><m:t>x</m:t></m:r></m:den></m:f>
    </m:oMath>"""
)
omml_limit_text = importers._omml_text(omml_limit)
if r"\lim_{x→0}" not in omml_limit_text or r"\frac{sin x}{x}" not in omml_limit_text:
    failures.append(f"OMML conversion: limit-low expression was not converted in {omml_limit_text!r}")

omml_matrix = ET.fromstring(
    """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:m>
        <m:mr><m:e><m:r><m:t>a</m:t></m:r></m:e><m:e><m:r><m:t>b</m:t></m:r></m:e></m:mr>
        <m:mr><m:e><m:r><m:t>c</m:t></m:r></m:e><m:e><m:r><m:t>d</m:t></m:r></m:e></m:mr>
      </m:m>
    </m:oMath>"""
)
if r"\begin{matrix}a & b \\ c & d\end{matrix}" not in importers._omml_text(omml_matrix):
    failures.append("OMML conversion: matrix was not converted")

equation = ET.fromstring('<hp:equation xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" script="x^2+1"/>')
if importers._hwpx_equation_text(equation) != "$x^2+1$":
    failures.append("HWPX equation conversion: script attr was not wrapped as editable math text")

# 1) DOCX import
result = importers.import_docx("sample.docx", make_docx(), {})
created = result["created"]
print("DOCX import:", len(created), "problems;", result["notices"])
if len(created) != 2:
    failures.append(f"DOCX: expected 2 problems, got {len(created)}")
if not created[0]["image_paths"]:
    failures.append("DOCX: first problem missing image")
if "광합성" not in created[0]["stem"]:
    failures.append("DOCX: stem text wrong")
if not created[0].get("tables"):
    failures.append("DOCX: table not captured into tables model")
elif "세포소기관" not in str(created[0]["tables"]):
    failures.append("DOCX: table content wrong")

omml_table_result = importers.import_docx("omml_table.docx", make_omml_table_docx(), {"tags": "omml-table"})
omml_table_created = omml_table_result["created"]
if not omml_table_created or r"\frac{1}{2}" not in str(omml_table_created[0].get("tables")):
    failures.append(f"DOCX import: table-cell OMML math lost {omml_table_created!r}")

# 2) image import
result = importers.import_image("pic.png", make_png(), {})
print("Image import:", len(result["created"]), "problems;", result["notices"])
image_paths_for_roundtrip = result["created"][0]["image_paths"] if result["created"] else []
if not image_paths_for_roundtrip:
    failures.append("Image: missing image path")

# 2b) HWP 편집본형 문항 경계: 문제 뒤의 [배점][출처 01] 라인으로도 분리되어야 한다.
trailer_chunks = importers._paragraphs_to_chunks(
    [
        ("의 값은?", [], []),
        ("[2점][2025학년도 수능 01]", [], []),
        ("함수 f(x)에 대하여 f(1)의 값은?", [], []),
        ("[3점][2025학년도 수능 02]", [], []),
    ]
)
trailer_sink = importers._create_from_chunks(trailer_chunks, "hwp", "trailer_math.hwp", {})
trailer_numbers = [item["number"] for item in trailer_sink.created]
print("HWP trailer split:", trailer_numbers)
if trailer_numbers != ["1", "2"]:
    failures.append(f"HWP trailer split: expected ['1', '2'], got {trailer_numbers}")

from app import docx_writer  # noqa: E402

# 2c) 레거시 보기/선지/주관식 괄호 템플릿
legacy_inline = "\n".join(
    [
        "12. 다음 <보기>를 보고 가장 적절한 것을 고르시오.",
        "<보기>",
        "ㄱ. 첫째 조건",
        "ㄴ. 둘째 조건",
        "</보기>",
        "1 사과 2 배 3 감 4 귤 5 포도",
    ]
)
legacy_stem, legacy_choices = importers._split_stem_and_choices(legacy_inline)
print("Legacy inline choices:", legacy_choices)
if legacy_choices != ["사과", "배", "감", "귤", "포도"]:
    failures.append(f"Legacy inline choices: wrong split {legacy_choices}")
if "<보기>" not in legacy_stem or "ㄱ. 첫째 조건" not in legacy_stem:
    failures.append("Legacy inline choices: 보기 text was not preserved in stem")

_, marker_only_choices = importers._split_stem_and_choices("① ② ③ ④ ⑤")
if marker_only_choices != ["①", "②", "③", "④", "⑤"]:
    failures.append(f"Legacy marker-only circled choices: wrong split {marker_only_choices}")
_, split_line_circled_choices = importers._split_stem_and_choices("① 120 ② 125 ③ 130\n④ 135 ⑤ 140")
if split_line_circled_choices != ["120", "125", "130", "135", "140"]:
    failures.append(f"Split-line circled choices: wrong split {split_line_circled_choices}")
leaked_choice_stem = "① a ② b ③ c\n④ d ⑤ e\nproblem body starts here"
cleaned_leaked_stem = importers._strip_leading_leaked_choice_block(
    leaked_choice_stem,
    ["current-a", "current-b", "current-c", "current-d", "current-e"],
)
if cleaned_leaked_stem != "problem body starts here":
    failures.append(f"HWP leaked choices: stale leading choices were not stripped: {cleaned_leaked_stem!r}")
choice_only_stem = "① a ② b ③ c ④ d ⑤ e"
preserved_choice_stem = importers._strip_leading_leaked_choice_block(
    choice_only_stem,
    ["current-a", "current-b", "current-c", "current-d", "current-e"],
)
if preserved_choice_stem != choice_only_stem:
    failures.append("HWP leaked choices: choice-only stem was stripped without a following body")
marker_repair_items = [
    {"number": "23", "unit": "[2점][2024년 5월 확통23]", "stem": "23"},
    {"number": "23", "unit": "[3점][2024년 5월 확통23]", "stem": "24"},
    {"number": "25", "unit": "[3점][2024년 5월 확통25]", "stem": "25"},
    {"number": "26", "unit": "[3점][2024년 5월 확통26]", "stem": "26"},
]
importers_hwp_ir._repair_problem_marker_sequence(marker_repair_items)
if [item["number"] for item in marker_repair_items] != ["23", "24", "25", "26"]:
    failures.append(f"HWP marker repair: wrong numbers {marker_repair_items}")
if marker_repair_items[1]["unit"] != "[3점][2024년 5월 확통24]":
    failures.append(f"HWP marker repair: wrong repaired unit {marker_repair_items[1]['unit']!r}")
_, numeric_marker_choices = importers._split_stem_and_choices("1 2 3 4 5")
if numeric_marker_choices != ["1", "2", "3", "4", "5"]:
    failures.append(f"Legacy marker-only numeric choices: wrong split {numeric_marker_choices}")

legacy_line_choices = "\n".join(
    [
        "13. 다음 보기에서 알 수 있는 내용은?",
        "보기",
        "자료의 핵심은 선택지와 구분되어야 한다.",
        "1. 사과",
        "2. 배",
        "3. 감",
        "4. 귤",
        "5. 포도",
    ]
)
line_stem, line_choices = importers._split_stem_and_choices(legacy_line_choices)
print("Legacy numbered line choices:", line_choices)
if line_choices != ["사과", "배", "감", "귤", "포도"]:
    failures.append(f"Legacy numbered line choices: wrong split {line_choices}")
if "자료의 핵심" not in line_stem:
    failures.append("Legacy numbered line choices: stem text lost")

numbered_view = "\n".join(
    [
        "14. 다음 <보기>의 내용을 확인하시오.",
        "<보기>",
        "1. 첫째 단서",
        "2. 둘째 단서",
        "</보기>",
    ]
)
view_stem, view_choices = importers._split_stem_and_choices(numbered_view)
if view_choices or "1. 첫째 단서" not in view_stem or "2. 둘째 단서" not in view_stem:
    failures.append("Legacy numbered view: non-choice numbered lines were misread as choices")

storage.ensure_dirs()
legacy_problem = {
    "number": "12",
    "title": "레거시 객관식",
    "stem": "12. 다음 <보기>를 보고 고르시오.\n<보기>\nㄱ. 첫째 조건\n</보기>",
    "choices": ["사과", "배", "감", "귤", "포도"],
    "answer": "1",
    "explanation": "",
    "image_paths": [],
    "tables": [],
}
legacy_hwpx_path = storage.EXPORT_DIR / "legacy_objective.hwpx"
legacy_docx_path = storage.EXPORT_DIR / "legacy_objective.docx"
hwpx_writer.write_hwpx(
    legacy_hwpx_path, "레거시 객관식", [legacy_problem], template_key="legacy_objective_12345"
)
docx_writer.write_docx(
    legacy_docx_path, "레거시 객관식", [legacy_problem], template_key="legacy_objective_12345"
)
legacy_hwpx_text = hwpx_text(legacy_hwpx_path)
legacy_docx_text = docx_text(legacy_docx_path)
if "1 사과" not in legacy_hwpx_text or "5 포도" not in legacy_hwpx_text:
    failures.append("Legacy objective HWPX: bare 1~5 choices missing")
if "1 사과" not in legacy_docx_text or "5 포도" not in legacy_docx_text:
    failures.append("Legacy objective DOCX: bare 1~5 choices missing")
if "① 사과" in legacy_hwpx_text or "1) 사과" in legacy_hwpx_text:
    failures.append("Legacy objective HWPX: wrong choice label style")

blank_problem = {
    "number": "15",
    "title": "레거시 주관식",
    "stem": "15. 다음 물음에 답하시오.",
    "choices": [],
    "answer": "42",
    "explanation": "",
    "image_paths": [],
    "tables": [],
}
blank_hwpx_path = storage.EXPORT_DIR / "legacy_blank.hwpx"
blank_docx_path = storage.EXPORT_DIR / "legacy_blank.docx"
hwpx_writer.write_hwpx(
    blank_hwpx_path, "레거시 주관식", [blank_problem], template_key="legacy_short_answer_blank"
)
docx_writer.write_docx(
    blank_docx_path, "레거시 주관식", [blank_problem], template_key="legacy_short_answer_blank"
)
if "(          )" not in hwpx_text(blank_hwpx_path):
    failures.append("Legacy blank HWPX: subjective answer blank missing")
if "(          )" not in docx_text(blank_docx_path):
    failures.append("Legacy blank DOCX: subjective answer blank missing")

# 2d) TXT/붙여넣기 입력: 공통 보기, 선지, 정답, 해설을 문항 필드로 분리
text_input = "\n".join(
    [
        "<보기>",
        "이 자료는 두 문항에 공통으로 쓰인다.",
        "</보기>",
        "1. 다음 중 알맞은 것은?",
        "1 사과 2 배 3 감 4 귤 5 포도",
        "정답: 2",
        "해설: 두 번째 선지가 정답이다.",
        "",
        "2. 다음 물음에 답하시오.",
        "답: 봄",
        "풀이: 계절을 묻는 주관식이다.",
    ]
)
text_result = importers.import_text("paste.txt", text_input.encode("utf-8"), {"subject": "국어"})
text_created = text_result["created"]
print("Text import:", len(text_created), "problems;", text_result["notices"])
if len(text_created) != 2:
    failures.append(f"Text import: expected 2 problems, got {len(text_created)}")
else:
    first, second = text_created
    if "<보기>" not in first["stem"] or "공통으로 쓰인다" not in first["stem"]:
        failures.append("Text import: common passage not attached to first question")
    if first["choices"] != ["사과", "배", "감", "귤", "포도"]:
        failures.append(f"Text import: choices wrong {first['choices']}")
    if first["answer"] != "2" or "두 번째" not in first["explanation"]:
        failures.append("Text import: answer/explanation not separated")
    if second["choices"]:
        failures.append("Text import: short answer problem should not have choices")
    if second["answer"] != "봄" or "주관식" not in second["explanation"]:
        failures.append("Text import: short answer/explanation not separated")

from fastapi.testclient import TestClient  # noqa: E402
from app.main import app  # noqa: E402

manual_client = TestClient(app)
manual_response = manual_client.post(
    "/api/import-text",
    json={
        "title": "직접 입력 테스트",
        "text": "31. 직접 입력 문항은?\n① 가능 ② 불가\n정답: 1",
        "metadata": {"tags": "manual-api"},
        "source_type": "manual",
    },
)
if manual_response.status_code != 200:
    failures.append(f"Manual import API: status {manual_response.status_code}")
else:
    manual_items = manual_response.json().get("created") or []
    if len(manual_items) != 1 or manual_items[0].get("source_type") != "manual":
        failures.append("Manual import API: did not create one manual problem")
    elif manual_items[0].get("choices") != ["가능", "불가"]:
        failures.append(f"Manual import API: choices wrong {manual_items[0].get('choices')}")

# 3) 내보내기 준비: 정답/선지/해설을 채운 뒤 in-memory로 문항을 잡아 둔다.
template = exam_templates.get_template(ROUNDTRIP_TEMPLATE_KEY)
if template.key != ROUNDTRIP_TEMPLATE_KEY or template.key == "basic":
    failures.append(f"HWPX template: expected non-default template {ROUNDTRIP_TEMPLATE_KEY!r}")
ids = [p["id"] for p in created]
for problem_id in ids:
    storage.update_problem(
        problem_id,
        {
            "answer": "2",
            "explanation": "엽록체에서 광합성이 일어난다.",
            "choices": ["미토콘드리아", "엽록체", "리보솜"],
        },
    )
# 표 출력 검증용: 첫 문항에 자료 표를 붙인다.
storage.update_problem(
    ids[0], {"tables": [[["구분", "결과물"], ["광합성", "포도당"], ["호흡", "에너지"]]]}
)
rich_problem = storage.create_problem(
    {
        "source_type": "manual",
        "source_name": "rich_roundtrip",
        "number": "3",
        "title": "수식과 긴 지문 유지",
        "stem": RICH_STEM,
        "choices": RICH_CHOICES,
        "answer": "3",
        "explanation": "수식, 긴 지문, 표, 이미지, 선지가 모두 유지되어야 한다.",
        "image_paths": image_paths_for_roundtrip,
        "tables": RICH_TABLES,
    }
)
ids.append(rich_problem["id"])
problems = storage.get_problems_by_ids(ids)
if not problems[0].get("tables"):
    failures.append("Tables: not persisted/round-tripped through storage")
rich_loaded = problems[-1]
if rich_loaded.get("choices") != RICH_CHOICES:
    failures.append("Rich content: choices not persisted as a separate list")
if not rich_loaded.get("image_paths"):
    failures.append("Rich content: image path not persisted")
if rich_loaded.get("tables") != RICH_TABLES:
    failures.append("Rich content: table not persisted")
if rich_loaded.get("math_summary", {}).get("count", 0) < len(RICH_FORMULAS):
    failures.append(f"Rich content: math summary too small {rich_loaded.get('math_summary')}")
if not any(span.get("field") == "tables" for span in rich_loaded.get("math_spans", [])):
    failures.append("Rich content: table math spans missing from API/storage model")
storage.ensure_dirs()

# 3a) 기본 HWPX 내보내기 → 출력 파일을 직접 검사(임포터/중복검사와 분리해 출력 충실도 검증)
export_path = storage.EXPORT_DIR / "roundtrip.hwpx"
hwpx_writer.write_hwpx(export_path, "라운드트립 테스트", problems, template_key=ROUNDTRIP_TEMPLATE_KEY)
body_text = hwpx_text(export_path)
print(f"HWPX export ({ROUNDTRIP_TEMPLATE_KEY} template): {len(body_text)} chars in body")
if "광합성" not in body_text or "세포 호흡" not in body_text:
    failures.append("HWPX export: stem text missing from section XML")

# 3a-2) 표(hp:tbl) 출력 검증: 구조 태그 + 셀 텍스트가 본문에 살아있는지
section_xml = hwpx_section_xml(export_path)
if "<hp:tbl" not in section_xml or "<hp:tc" not in section_xml:
    failures.append("HWPX export: hp:tbl/hp:tc structure missing")
if "구분" not in body_text or "포도당" not in body_text:
    failures.append("HWPX export: table cell text missing from section XML")
missing_rich_formulas = [formula for formula in RICH_FORMULAS if formula not in body_text]
if missing_rich_formulas:
    failures.append(f"HWPX export: formula text missing: {missing_rich_formulas}")
if 'charPrIDRef="5"' not in section_xml:
    failures.append("HWPX export: math text was not emitted as separate math charPr runs")
if not re.search(r"<hp:tc[\s\S]*?charPrIDRef=\"5\"[\s\S]*?\\frac\{1\}\{2\}[\s\S]*?</hp:tc>", section_xml):
    failures.append("HWPX export: table-cell math was not emitted as math charPr runs")
if "<hp:equation" in section_xml:
    failures.append("HWPX export: native equations should stay off on the default compatibility path")

native_math_path = storage.EXPORT_DIR / "roundtrip_native_math.hwpx"
hwpx_writer.write_hwpx(
    native_math_path,
    "네이티브 수식 테스트",
    [rich_loaded],
    template_key=ROUNDTRIP_TEMPLATE_KEY,
    native_math=True,
)
native_section_xml = hwpx_section_xml(native_math_path)
native_section_text = html.unescape(native_section_xml)
if "<hp:equation" not in native_section_xml or "<hp:script>" not in native_section_xml:
    failures.append("HWPX native math: expected hp:equation script objects were not emitted")
for script in (
    "{1} over {2}",
    "sqrt {x}",
    "sum",
    "int",
    "bar {x}",
    "vec {v}",
    "matrix{a & b # c & d}",
    "(x+1)",
    "^3sqrt {x^{2}+1}",
    "{n} choose {r}",
    "x∈ℝ",
    "A∪ B",
    "vec {AB}",
    "hat {ABC}",
):
    if script not in native_section_text:
        failures.append(f"HWPX native math: expected HancomEQN script {script!r} missing")
native_import = importers.import_hwpx("roundtrip_native_math.hwpx", native_math_path.read_bytes(), {})
native_text = "\n".join(p.get("stem", "") for p in native_import["created"])
if "$" not in native_text or "over" not in native_text:
    failures.append(f"HWPX native math import: equation script was not preserved as editable math text {native_text!r}")
native_tables_text = "\n".join(str(p.get("tables", "")) for p in native_import["created"])
if "over" not in native_tables_text or "sqrt" not in native_tables_text or "sum" not in native_tables_text:
    failures.append(f"HWPX native math import: table-cell equation scripts were not preserved {native_tables_text!r}")

fallback_v2_path = storage.EXPORT_DIR / "fallback_labels_v2.hwpx"
hwpx_writer_v2.write_hwpx(
    fallback_v2_path,
    "fallback label test",
    [
        {
            "source_type": "manual",
            "source_name": "fallback",
            "number": "1",
            "title": "",
            "stem": "",
            "choices": [],
            "answer": "",
            "explanation": "",
            "image_paths": [],
            "tables": [],
        }
    ],
    template_key="kice_math",
    native_math=True,
)
fallback_v2_text = hwpx_text(fallback_v2_path)
if "1. 문제" not in fallback_v2_text:
    failures.append(f"HWPX v2 fallback label: expected Korean fallback heading, got {fallback_v2_text!r}")
if "臾몄젣" in fallback_v2_text:
    failures.append("HWPX v2 fallback label: mojibake text leaked into output")

from app import main as app_main  # noqa: E402

api_default_response = app_main.export(
    app_main.ExportPayload(
        ids=[rich_loaded["id"]],
        title="API 기본 수식",
        format="hwpx",
        template_key=ROUNDTRIP_TEMPLATE_KEY,
        native_math=False,
    )
)
api_default_export_path = Path(str(api_default_response.path))
api_default_section_xml = hwpx_section_xml(api_default_export_path)
if "<hp:equation" in api_default_section_xml:
    failures.append("API export: default HWPX path should not emit native equations")
if not hwpx_has_distinct_math_run(api_default_export_path, r"\frac{1}{2}", in_table=False):
    failures.append("API export: default HWPX path did not preserve math styled runs")
if not hwpx_has_distinct_math_run(api_default_export_path, r"\frac{1}{2}", in_table=True):
    failures.append("API export: default HWPX path did not preserve table-cell math styled runs")

api_response = app_main.export(
    app_main.ExportPayload(
        ids=[rich_loaded["id"]],
        title="API 네이티브 수식",
        format="hwpx",
        template_key=ROUNDTRIP_TEMPLATE_KEY,
        native_math=True,
    )
)
api_export_path = Path(str(api_response.path))
api_section_xml = hwpx_section_xml(api_export_path)
if "<hp:equation" not in api_section_xml:
    failures.append("API export: native_math payload did not route to native HWPX equations")
api_template_default_response = app_main.export(
    app_main.ExportPayload(
        ids=[rich_loaded["id"]],
        title="API template native math default",
        format="hwpx",
        template_key="kice_math",
    )
)
api_template_default_export_path = Path(str(api_template_default_response.path))
api_template_default_section_xml = hwpx_section_xml(api_template_default_export_path)
if "<hp:equation" not in api_template_default_section_xml:
    failures.append("API export: kice_math should default to native HWPX equations when native_math is omitted")
api_explicit_off_response = app_main.export(
    app_main.ExportPayload(
        ids=[rich_loaded["id"]],
        title="API explicit native math off",
        format="hwpx",
        template_key="kice_math",
        native_math=False,
    )
)
api_explicit_off_export_path = Path(str(api_explicit_off_response.path))
api_explicit_off_section_xml = hwpx_section_xml(api_explicit_off_export_path)
if "<hp:equation" in api_explicit_off_section_xml:
    failures.append("API export: explicit native_math=false should override kice_math template default")
preview_calls: list[dict[str, Any]] = []
original_preview_available = app_main.preview.available
original_preview_render = app_main.preview.render_preview
try:
    app_main.preview.available = lambda: True

    def fake_render_preview(title, preview_problems, template_key="basic", **kwargs):
        preview_calls.append(
            {
                "title": title,
                "problem_count": len(preview_problems),
                "template_key": template_key,
                "native_math": kwargs.get("native_math"),
                "include_answer_sheet": kwargs.get("include_answer_sheet"),
            }
        )
        return {"pages": [], "page_count": 1, "truncated": False}

    app_main.preview.render_preview = fake_render_preview
    app_main.preview_export(
        app_main.ExportPayload(
            ids=[rich_loaded["id"]],
            title="Preview template native math default",
            format="hwpx",
            template_key="kice_math",
        )
    )
    app_main.preview_export(
        app_main.ExportPayload(
            ids=[rich_loaded["id"]],
            title="Preview explicit native math off",
            format="hwpx",
            template_key="kice_math",
            native_math=False,
        )
    )
    app_main.preview_export(
        app_main.ExportPayload(
            ids=[rich_loaded["id"]],
            title="Preview docx native math off",
            format="docx",
            template_key="kice_math",
        )
    )
finally:
    app_main.preview.available = original_preview_available
    app_main.preview.render_preview = original_preview_render
expected_preview_native_math = [True, False, False]
actual_preview_native_math = [call.get("native_math") for call in preview_calls]
if actual_preview_native_math != expected_preview_native_math:
    failures.append(
        "API preview: native_math effective values do not match export rules "
        f"{actual_preview_native_math}"
    )
if any(call.get("problem_count") != 1 for call in preview_calls):
    failures.append(f"API preview: selected problems were not loaded correctly {preview_calls}")
if "정서의 변화와 서술자의 태도" not in body_text:
    failures.append("HWPX export: long Korean passage missing from section XML")
if "유지 확인" not in body_text or "정서 변화" not in body_text:
    failures.append("HWPX export: rich table cell text missing from section XML")
if section_xml.count("<hc:img") < 2:
    failures.append("HWPX export: expected images for imported and rich problems")
body_paragraphs = hwpx_paragraph_texts(export_path)
formatted_rich_choices = [
    f"{label} {choice}" for label, choice in zip(RICH_CHOICE_LABELS, RICH_CHOICES)
]
missing_choice_paragraphs = [
    choice for choice in formatted_rich_choices if choice not in body_paragraphs
]
if missing_choice_paragraphs:
    failures.append(
        f"HWPX export: rich choices not emitted as separate paragraphs: {missing_choice_paragraphs}"
    )
print("Rich HWPX export: formulas, passage, table, image, choices checked")

# 3b) rhwp 엔진으로 생성 HWPX 구조 검증 + 렌더링 (설치된 경우)
try:
    import rhwp
except Exception:
    rhwp = None
if rhwp is not None:
    doc = rhwp.parse(str(export_path))
    rendered = bytes(doc.render_png(0))
    print(f"rhwp validation: pages={doc.page_count}, render={len(rendered)} bytes")
    if doc.page_count < 1 or not rendered:
        failures.append("rhwp: generated HWPX failed to parse/render")
    if "광합성" not in doc.extract_text():
        failures.append("rhwp: text not extractable from generated HWPX")
    native_doc = rhwp.parse(str(native_math_path))
    native_rendered = bytes(native_doc.render_png(0))
    print(f"rhwp native math validation: pages={native_doc.page_count}, render={len(native_rendered)} bytes")
    if native_doc.page_count < 1 or not native_rendered:
        failures.append("rhwp: native math HWPX failed to parse/render")
    api_default_doc = rhwp.parse(str(api_default_export_path))
    api_default_rendered = bytes(api_default_doc.render_png(0))
    print(
        "rhwp API default validation: "
        f"pages={api_default_doc.page_count}, render={len(api_default_rendered)} bytes"
    )
    if api_default_doc.page_count < 1 or not api_default_rendered:
        failures.append("rhwp: API default HWPX failed to parse/render")
    api_native_doc = rhwp.parse(str(api_export_path))
    api_native_rendered = bytes(api_native_doc.render_png(0))
    print(
        "rhwp API native math validation: "
        f"pages={api_native_doc.page_count}, render={len(api_native_rendered)} bytes"
    )
    if api_native_doc.page_count < 1 or not api_native_rendered:
        failures.append("rhwp: API native math HWPX failed to parse/render")
else:
    print("rhwp not installed; skipping render validation")

# 3c) 정답·해설지 포함 HWPX — 출력 파일 직접 검사 + 페이지 나눔 확인
sheet_path = storage.EXPORT_DIR / "answer_sheet.hwpx"
hwpx_writer.write_hwpx(
    sheet_path, "정답지 테스트", problems, template_key=ROUNDTRIP_TEMPLATE_KEY, include_answer_sheet=True
)
sheet_text = hwpx_text(sheet_path)
if "빠른 정답" not in sheet_text or "②" not in sheet_text:
    failures.append("Answer sheet: quick answers missing from exported HWPX")
if "엽록체에서 광합성이 일어난다" not in sheet_text:
    failures.append("Answer sheet: explanation missing from exported HWPX")
if rhwp is not None:
    sheet_doc = rhwp.parse(str(sheet_path))
    print(f"Answer sheet pages: {sheet_doc.page_count}")
    if sheet_doc.page_count < 2:
        failures.append("Answer sheet: expected page break to add a page")

# 3d) DOCX 정답·해설지 — 출력 파일 직접 검사
docx_sheet_path = storage.EXPORT_DIR / "answer_sheet.docx"
docx_writer.write_docx(
    docx_sheet_path, "정답지 테스트", problems, template_key=ROUNDTRIP_TEMPLATE_KEY, include_answer_sheet=True
)
docx_body = docx_text(docx_sheet_path)
docx_xml = docx_document_xml(docx_sheet_path)
docx_math_text = "\n".join(docx_omml_texts(docx_sheet_path))
if "빠른 정답" not in docx_body or "②" not in docx_body:
    failures.append("Answer sheet: quick answers missing from exported DOCX")
if (
    "<m:oMath" not in docx_xml
    or "<m:f>" not in docx_xml
    or "<m:rad>" not in docx_xml
    or "<m:nary>" not in docx_xml
    or "<m:acc>" not in docx_xml
    or "<m:d>" not in docx_xml
    or "<m:m>" not in docx_xml
):
    failures.append("DOCX export: expected native OMML math objects were not emitted")
for formula in (
    r"\frac{1}{2}",
    r"\frac{\sqrt{x}}{2}",
    r"\sum_{k=1}^{n}",
    r"\sum_{k=1}^{n} k",
    r"\int_{0}^{1} dx",
    r"\overline{x}",
    r"\vec{v}",
    r"\left(x+1\right)",
    r"\begin{matrix}a & b \\ c & d\end{matrix}",
):
    if formula not in docx_math_text:
        failures.append(f"DOCX export: OMML roundtrip text missing {formula!r} in {docx_math_text!r}")
if not re.search(r"<w:tc[\s\S]*?<m:f>[\s\S]*?<m:t>1</m:t>[\s\S]*?<m:t>2</m:t>[\s\S]*?</w:tc>", docx_xml):
    failures.append("DOCX export: table-cell math was not emitted as native OMML")
print("DOCX answer sheet: ok" if "빠른 정답" in docx_body else "DOCX answer sheet: MISSING")

unicode_docx_path = storage.EXPORT_DIR / "unicode_math.docx"
docx_writer.write_docx(
    unicode_docx_path,
    "유니코드 수식 DOCX",
    [
        {
            "source_type": "manual",
            "source_name": "unicode_math",
            "title": "유니코드 수식",
            "stem": "√x, √(x+1), ∑_{k=1}^{n} k, ∫_0^1 f(x)dx",
            "choices": [],
            "answer": "",
            "explanation": "",
            "image_paths": [],
            "tables": [],
        }
    ],
)
unicode_docx_xml = docx_document_xml(unicode_docx_path)
unicode_docx_math_text = "\n".join(docx_omml_texts(unicode_docx_path))
if unicode_docx_xml.count("<m:rad>") < 2 or unicode_docx_xml.count("<m:nary>") < 2:
    failures.append("DOCX export: unicode radical/n-ary input was not emitted as native OMML")
for formula in (r"\sqrt{x}", r"\sqrt{(x+1)}", r"\sum_{k=1}^{n} k", r"\int_{0}^{1} f(x)dx"):
    if formula not in unicode_docx_math_text:
        failures.append(f"DOCX export: unicode OMML roundtrip text missing {formula!r} in {unicode_docx_math_text!r}")

# 3e) HWPX 임포터 라운드트립: 빈 DB에서 다시 읽어 본문이 살아있는지 확인
reset_problems()
result = importers.import_hwpx("roundtrip.hwpx", export_path.read_bytes(), {})
roundtrip_text = "\n".join(p["stem"] for p in result["created"])
print("HWPX importer roundtrip:", len(result["created"]), "problems")
if "광합성" not in roundtrip_text or "세포 호흡" not in roundtrip_text:
    failures.append("HWPX roundtrip: text lost on re-import")
roundtrip_tables = [p.get("tables") for p in result["created"] if p.get("tables")]
if not roundtrip_tables:
    failures.append("HWPX roundtrip: table lost on re-import")
elif "포도당" not in str(roundtrip_tables):
    failures.append("HWPX roundtrip: table cell text lost on re-import")
rich_roundtrip = next(
    (p for p in result["created"] if "x^2" in p["stem"] and r"\frac{1}{2}" in p["stem"]),
    None,
)
if rich_roundtrip is None:
    failures.append("HWPX roundtrip: rich formula/passage problem missing on re-import")
else:
    rich_text = rich_roundtrip["stem"]
    missing_rich_roundtrip_formulas = [
        formula for formula in RICH_FORMULAS if formula not in rich_text
    ]
    if missing_rich_roundtrip_formulas:
        failures.append(
            f"HWPX roundtrip: formula text lost on re-import: {missing_rich_roundtrip_formulas}"
        )
    if "정서의 변화와 서술자의 태도" not in rich_text:
        failures.append("HWPX roundtrip: long Korean passage lost on re-import")
    if not rich_roundtrip.get("image_paths"):
        failures.append("HWPX roundtrip: rich problem image lost on re-import")
    rich_roundtrip_tables = rich_roundtrip.get("tables") or []
    if "정서 변화" not in str(rich_roundtrip_tables):
        failures.append("HWPX roundtrip: rich problem table lost on re-import")
    roundtrip_choice_lines = [
        line.strip()
        for line in rich_text.splitlines()
        if line.strip().startswith(RICH_CHOICE_LABELS)
    ]
    choices_as_list = rich_roundtrip.get("choices") == RICH_CHOICES
    choices_as_lines = roundtrip_choice_lines[-len(RICH_CHOICES) :] == formatted_rich_choices
    if not choices_as_list and not choices_as_lines:
        failures.append("HWPX roundtrip: rich choices lost their separation on re-import")
print("Rich HWPX roundtrip: formulas, passage, table, image, choices checked")

# 4) CSV import still fine
csv_data = "번호,문제,정답\n1,사과는 영어로?,apple\n2,바다는 영어로?,sea\n".encode("utf-8-sig")
result = importers.import_csv("words.csv", csv_data, {})
print("CSV import:", len(result["created"]), "problems")
if len(result["created"]) != 2:
    failures.append("CSV: expected 2 problems")

# 5) 중복 감지: 같은 CSV를 다시 가져오면 0개 생성 + 건너뜀 안내
result = importers.import_csv("words.csv", csv_data, {})
print("CSV re-import (dedup):", len(result["created"]), "created;", result["notices"])
if result["created"]:
    failures.append("Dedup: re-import should create 0 problems")
if not any("건너뛰" in notice for notice in result["notices"]):
    failures.append("Dedup: missing skip notice on re-import")

# 5b) 한 파일 안의 동일 문항도 한 번만 생성
dup_csv = "번호,문제,정답\n1,중복질문,A\n2,중복질문,A\n3,다른질문,B\n".encode("utf-8-sig")
result = importers.import_csv("dups.csv", dup_csv, {})
print("CSV intra-file dedup:", len(result["created"]), "created")
if len(result["created"]) != 2:
    failures.append(f"Dedup: intra-file expected 2 unique, got {len(result['created'])}")

# 5c) HWP IR marker/choice/image sync regression.
original_ordered_stream = importers_hwp_ir._ordered_stream
try:
    marker13 = "[4\uc810][2025\ud559\ub144\ub3c4 \uc218\ub2a5 13]"
    marker14 = "[4\uc810][2025\ud559\ub144\ub3c4 \uc218\ub2a5 14]"
    marker30 = "[4\uc810][2025\ud559\ub144\ub3c4 \uc218\ub2a5 \uae30\ud55830]"
    circled = "\u2460 A\t\u2461 B\t\u2462 C\t\u2463 D\t\u2464 E"

    def fake_ordered_stream(payload, filename, save_image):
        return [
            ("text", "problem 13 stem"),
            ("text", marker13),
            ("text", circled),
            ("text", ""),
            ("image", "img13.png"),
            ("text", "problem 14 stem"),
            ("text", marker14),
            ("text", ""),
            ("image", "img14.png"),
            ("text", "\u2460 1\t\u2461 2\t\u2462 3"),
            ("text", "\u2463 4\t\u2464 5"),
            ("text", "problem 30 stem"),
            ("text", marker30),
            ("image", "img30.png"),
            ("table", [["2025\ud559\ub144\ub3c4 \uc218\ub2a5 \uc218\ud559(\ube60\ub978 \uc815\ub2f5)"]]),
            ("text", "[\uc815\ub2f5] 1"),
        ]

    importers_hwp_ir._ordered_stream = fake_ordered_stream
    synced = importers_hwp_ir.hwp_to_problems(
        b"",
        "sync.hwp",
        lambda name, data: name,
        importers._split_inline_circled_choices,
        split_stem_choices=importers._split_stem_and_choices,
    )
finally:
    importers_hwp_ir._ordered_stream = original_ordered_stream

print("HWP IR sync regression:", len(synced or []), "problems")
if not synced or len(synced) != 3:
    failures.append(f"HWP IR sync: expected 3 problems, got {len(synced or [])}")
else:
    if synced[0].get("stem") != "problem 13 stem" or synced[0].get("image_paths") != ["img13.png"]:
        failures.append(f"HWP IR sync: trailing image did not stay with problem 13: {synced[0]}")
    if synced[1].get("stem") != "problem 14 stem" or synced[1].get("image_paths") != ["img14.png"]:
        failures.append(f"HWP IR sync: marker-blank image did not stay with problem 14: {synced[1]}")
    if synced[2].get("stem") != "problem 30 stem" or synced[2].get("image_paths") != ["img30.png"]:
        failures.append(f"HWP IR sync: final problem image/cutoff failed: {synced[2]}")
    if any("\uc815\ub2f5" in str(problem.get("tables") or "") for problem in synced):
        failures.append(f"HWP IR sync: answer section leaked into problems: {synced}")

print()
if failures:
    print("FAILURES:")
    for failure in failures:
        print(" -", failure)
    sys.exit(1)
print("ALL OK")
